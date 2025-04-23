"""
Approval Authority UI module.
Handles the UI components for the approval authority role.
"""

import streamlit as st
import pandas as pd
from datetime import datetime
import plotly.express as px
import logging

from app.ui.common import (
    display_header,
    display_info_box,
    display_success_box,
    display_warning_box,
    display_error_box,
    show_loading_spinner
)
from app.database.queries import (
    get_pending_requests,
    get_request_by_id,
    update_request_status,
    get_document,
    search_requests,
    calculate_remaining_budget,
    get_top_destinations,
    get_faculty_travel_frequency
)
from app.services.document_service import extract_text_from_file
from app.services.ai_service import generate_ai_notes

def show_approval_dashboard():
    """Approval authority main dashboard"""
    display_header("Approval Authority Dashboard")
    
    # Show tabs for different approval authority functions
    tabs = st.tabs([
        "Pending Requests", 
        "Search & View", 
        "Analytics"
    ])
    
    # Tab 1: Pending Requests
    with tabs[0]:
        show_pending_requests()
    
    # Tab 2: Search & View
    with tabs[1]:
        show_search_interface()
    
    # Tab 3: Analytics
    with tabs[2]:
        show_approval_analytics()

def show_pending_requests():
    """Display pending requests for approval"""
    st.subheader("Pending Requests")
    
    # Get budget information for context
    total_budget, total_expenses, remaining_budget = calculate_remaining_budget()
    
    # Display budget context
    if remaining_budget <= 0:
        display_warning_box(
            "⚠️ The conference travel budget has been fully allocated. "
            "Please consider this when reviewing new requests."
        )
    else:
        display_info_box(
            f"Current remaining budget: ${remaining_budget:,.2f} "
            f"(Used: ${total_expenses:,.2f} of ${total_budget:,.2f})"
        )
    
    # Get pending requests
    with show_loading_spinner("Loading pending requests..."):
        pending_requests = get_pending_requests()
    
    if not pending_requests:
        display_info_box("There are no pending requests to review.")
        return
    
    # Prepare data for display
    requests_df = pd.DataFrame(pending_requests)
    
    # Format dates for display
    requests_df['created_at'] = pd.to_datetime(requests_df['created_at']).dt.strftime('%Y-%m-%d')
    requests_df['date_from'] = pd.to_datetime(requests_df['date_from']).dt.strftime('%Y-%m-%d')
    requests_df['date_to'] = pd.to_datetime(requests_df['date_to']).dt.strftime('%Y-%m-%d')
    
    # Add total cost column
    requests_df['total_cost'] = requests_df['per_diem'] + requests_df['registration_fee'] + requests_df['visa_fee']
    
    # Select columns for display
    display_df = requests_df[[
        'request_id', 'faculty_name', 'department', 'conference_name',
        'destination', 'date_from', 'date_to', 'total_cost', 'created_at'
    ]]
    
    # Rename columns for better display
    display_df.columns = [
        'Request ID', 'Faculty', 'Department', 'Conference',
        'Destination', 'Start Date', 'End Date', 'Total Cost (USD)', 'Submitted On'
    ]
    
    # Display the dataframe with a callback for selection
    selected_indices = st.dataframe(
        display_df,
        use_container_width=True,
        column_config={
            "Total Cost (USD)": st.column_config.NumberColumn(
                "Total Cost (USD)",
                format="$%.2f",
            ),
        }
    )
    
    # Let user select a request to review
    selected_request_id = st.selectbox(
        "Select a request to review",
        options=requests_df['request_id'].tolist(),
        format_func=lambda x: f"{x} - {requests_df[requests_df['request_id'] == x]['faculty_name'].values[0]} - {requests_df[requests_df['request_id'] == x]['conference_name'].values[0]}"
    )
    
    if selected_request_id:
        show_request_review(selected_request_id)

def show_request_review(request_id):
    """Display detailed request information for review"""
    st.subheader("Request Review")
    
    # Get request details
    request = get_request_by_id(request_id)
    
    if not request:
        st.error("Request not found.")
        return
    
    # Display request overview
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown(f"**Faculty:** {request['faculty_name']}")
        st.markdown(f"**Department:** {request['department']}")
        st.markdown(f"**Conference:** {request['conference_name']}")
        st.markdown(f"**Index Type:** {request['index_type']}")
        st.markdown(f"**Purpose:** {request['purpose_of_attending']}")
    
    with col2:
        st.markdown(f"**Destination:** {request['destination']}, {request['city']}")
        st.markdown(f"**Dates:** {request['date_from']} to {request['date_to']}")
        
        # Calculate total cost
        total_cost = request['per_diem'] + request['registration_fee'] + request['visa_fee']
        
        st.markdown(f"**Per Diem:** ${request['per_diem']}")
        st.markdown(f"**Registration Fee:** ${request['registration_fee']}")
        st.markdown(f"**Visa Fee:** ${request['visa_fee']}")
        st.markdown(f"**Total Cost:** ${total_cost}")
    
    # Display AI-generated summaries
    st.subheader("AI Analysis")
    
    summary_tabs = st.tabs([
        "Conference Summary", 
        "Research Summary", 
        "Travel Notes"
    ])
    
    with summary_tabs[0]:
        st.markdown(request['conference_summary'])
    
    with summary_tabs[1]:
        st.markdown(request['research_summary'])
    
    with summary_tabs[2]:
        st.markdown(request['notes_summary'])
    
    # Display document viewers
    st.subheader("Attached Documents")
    
    doc_tabs = st.tabs([
        "Research Paper", 
        "Acceptance Letter"
    ])
    
    with doc_tabs[0]:
        research_paper = get_document(request_id, "research_paper")
        if research_paper:
            if st.button("View Research Paper"):
                try:
                    from io import BytesIO
                    
                    # Extract text from document
                    import PyPDF2
                    import docx
                    
                    if research_paper['file_name'].endswith('.pdf'):
                        pdf_reader = PyPDF2.PdfReader(BytesIO(research_paper['file_data']))
                        text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                    elif research_paper['file_name'].endswith('.docx'):
                        doc = docx.Document(BytesIO(research_paper['file_data']))
                        text = "\n".join([para.text for para in doc.paragraphs])
                    else:
                        text = "Unsupported file format"
                    
                    st.text_area("Research Paper Content", text, height=300)
                except Exception as e:
                    st.error(f"Error displaying document: {str(e)}")
        else:
            st.info("No research paper attached.")
    
    with doc_tabs[1]:
        acceptance_letter = get_document(request_id, "acceptance_letter")
        if acceptance_letter:
            if st.button("View Acceptance Letter"):
                try:
                    from io import BytesIO
                    
                    # Extract text from document
                    import PyPDF2
                    import docx
                    
                    if acceptance_letter['file_name'].endswith('.pdf'):
                        pdf_reader = PyPDF2.PdfReader(BytesIO(acceptance_letter['file_data']))
                        text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                    elif acceptance_letter['file_name'].endswith('.docx'):
                        doc = docx.Document(BytesIO(acceptance_letter['file_data']))
                        text = "\n".join([para.text for para in doc.paragraphs])
                    else:
                        text = "Unsupported file format"
                    
                    st.text_area("Acceptance Letter Content", text, height=300)
                except Exception as e:
                    st.error(f"Error displaying document: {str(e)}")
        else:
            st.info("No acceptance letter attached.")
    
    # Budget impact analysis
    st.subheader("Budget Impact")
    
    total_budget, total_expenses, remaining_budget = calculate_remaining_budget()
    
    # Calculate percentage of remaining budget
    budget_percentage = (total_cost / remaining_budget) * 100 if remaining_budget > 0 else 100
    
    # Display budget impact
    if budget_percentage > 100:
        display_error_box(
            f"⚠️ This request exceeds the remaining budget by ${total_cost - remaining_budget:,.2f}. "
            "Approving will require a budget adjustment."
        )
    elif budget_percentage > 50:
        display_warning_box(
            f"⚠️ This request will use {budget_percentage:.1f}% of the remaining budget. "
            f"Remaining after approval: ${remaining_budget - total_cost:,.2f}"
        )
    else:
        display_info_box(
            f"This request will use {budget_percentage:.1f}% of the remaining budget. "
            f"Remaining after approval: ${remaining_budget - total_cost:,.2f}"
        )
    
    # Decision making form
    st.subheader("Make Decision")
    
    with st.form("decision_form"):
        decision = st.radio(
            "Decision",
            options=["Approve", "Reject"],
            horizontal=True
        )
        
        # Option to generate AI recommendation
        if st.checkbox("Generate AI recommendation"):
            with show_loading_spinner("Generating recommendation..."):
                ai_recommendation = generate_ai_notes(
                    request['conference_name'],
                    request['purpose_of_attending'], 
                    request['index_type'],
                    request['destination'], 
                    request['city']
                )
                st.info(ai_recommendation)
        
        notes = st.text_area(
            "Decision Notes",
            placeholder="Provide notes explaining your decision..."
        )
        
        # Submit button
        submitted = st.form_submit_button("Submit Decision")
        
        if submitted:
            if not notes:
                st.error("Please provide decision notes before submitting.")
            else:
                try:
                    status = "approved" if decision == "Approve" else "rejected"
                    
                    # Check budget before approving
                    if status == "approved" and budget_percentage > 100:
                        st.error(
                            "Cannot approve: This request exceeds the remaining budget. "
                            "Please adjust the budget before approving."
                        )
                    else:
                        # Update request status
                        update_request_status(request_id, status, notes)
                        
                        if status == "approved":
                            display_success_box("Request approved successfully!")
                        else:
                            display_info_box("Request rejected.")
                        
                        # Refresh the page to show updated pending requests
                        st.experimental_rerun()
                        
                except Exception as e:
                    st.error(f"Error processing decision: {str(e)}")

def show_search_interface():
    """Display interface for searching and viewing requests"""
    st.subheader("Search & View Requests")
    
    # Search filters
    with st.expander("Search Filters", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            status = st.selectbox(
                "Status",
                options=["All", "Pending", "Approved", "Rejected"]
            )
            
            faculty_name = st.text_input(
                "Faculty Name",
                placeholder="Enter faculty name..."
            )
            
            department = st.selectbox(
                "Department",
                options=["All", "Computer Science", "Engineering", "Business", "Mathematics", "Physics"]
            )
        
        with col2:
            destination = st.text_input(
                "Destination",
                placeholder="Enter country or city..."
            )
            
            date_range = st.date_input(
                "Date Range",
                value=[],
                help="Select date range for travel"
            )
        
        # Prepare filters
        filters = {}
        
        if status != "All":
            filters["status"] = status.lower()
            
        if faculty_name:
            filters["faculty_name"] = faculty_name
            
        if department != "All":
            filters["department"] = department
            
        if destination:
            filters["destination"] = destination
            
        if len(date_range) == 2:
            filters["date_from"] = date_range[0]
            filters["date_to"] = date_range[1]
        
        # Search button
        search_clicked = st.button("Search")
        
        if search_clicked:
            # Execute search
            with show_loading_spinner("Searching..."):
                search_results = search_requests(filters)
                st.session_state.search_results = search_results
    
    # Display search results
    if "search_results" in st.session_state:
        results = st.session_state.search_results
        
        if not results:
            st.info("No requests found matching your search criteria.")
        else:
            # Convert to DataFrame for display
            df = pd.DataFrame(results)
            
            # Format dates
            df['created_at'] = pd.to_datetime(df['created_at']).dt.strftime('%Y-%m-%d')
            df['date_from'] = pd.to_datetime(df['date_from']).dt.strftime('%Y-%m-%d')
            df['date_to'] = pd.to_datetime(df['date_to']).dt.strftime('%Y-%m-%d')
            
            # Add total cost column
            df['total_cost'] = df['per_diem'] + df['registration_fee'] + df['visa_fee']
            
            # Select columns for display
            display_df = df[[
                'request_id', 'faculty_name', 'department', 'conference_name',
                'destination', 'date_from', 'date_to', 'total_cost', 'status', 'created_at'
            ]]
            
            # Rename columns for better display
            display_df.columns = [
                'Request ID', 'Faculty', 'Department', 'Conference',
                'Destination', 'Start Date', 'End Date', 'Total Cost (USD)', 'Status', 'Submitted On'
            ]
            
            # Status color function
            def status_color(val):
                if val == 'approved':
                    return 'background-color: #d4edda; color: #155724'
                elif val == 'rejected':
                    return 'background-color: #f8d7da; color: #721c24'
                elif val == 'pending':
                    return 'background-color: #fff3cd; color: #856404'
                return ''
            
            # Apply styling
            styled_df = display_df.style.applymap(status_color, subset=['Status'])
            
            # Display results
            st.dataframe(
                styled_df,
                use_container_width=True,
                column_config={
                    "Total Cost (USD)": st.column_config.NumberColumn(
                        "Total Cost (USD)",
                        format="$%.2f",
                    ),
                }
            )
            
            # View details
            selected_id = st.selectbox(
                "Select a request to view details",
                options=df['request_id'].tolist(),
                format_func=lambda x: f"{x} - {df[df['request_id'] == x]['faculty_name'].values[0]} - {df[df['request_id'] == x]['conference_name'].values[0]}"
            )
            
            if selected_id:
                show_request_review(selected_id)

def show_approval_analytics():
    """Display analytics for approval decisions"""
    st.subheader("Approval Analytics")
    
    # Budget overview
    total_budget, total_expenses, remaining_budget = calculate_remaining_budget()
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric(
            "Total Budget",
            f"${total_budget:,.2f}"
        )
    
    with col2:
        st.metric(
            "Total Spent",
            f"${total_expenses:,.2f}"
        )
    
    with col3:
        st.metric(
            "Remaining Budget",
            f"${remaining_budget:,.2f}"
        )
    
    # Budget utilization chart
    st.subheader("Budget Utilization")
    
    budget_data = pd.DataFrame({
        "Category": ["Spent", "Remaining"],
        "Amount": [total_expenses, remaining_budget]
    })
    
    if not budget_data.empty and total_budget > 0:
        fig = px.pie(
            budget_data,
            values="Amount",
            names="Category",
            title="Budget Utilization",
            color="Category",
            color_discrete_map={"Spent": "#4285F4", "Remaining": "#34A853"}
        )
        
        st.plotly_chart(fig, use_container_width=True)
    
    # Top destinations
    st.subheader("Top Travel Destinations")
    
    top_destinations = get_top_destinations()
    
    if top_destinations:
        dest_df = pd.DataFrame(top_destinations)
        
        fig = px.bar(
            dest_df,
            x="destination",
            y="count",
            title="Most Frequent Travel Destinations",
            labels={"destination": "Destination", "count": "Number of Trips"}
        )
        
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("No destination data available.")
    
    # Faculty travel frequency
    st.subheader("Faculty Travel Frequency")
    
    faculty_travel = get_faculty_travel_frequency()
    
    if faculty_travel:
        faculty_df = pd.DataFrame(faculty_travel)
        
        # Format for display
        faculty_df["total_cost"] = faculty_df["total_cost"].apply(lambda x: f"${x:,.2f}")
        
        st.dataframe(
            faculty_df,
            column_config={
                "faculty_name": "Faculty Name",
                "department": "Department",
                "travel_count": "Number of Trips",
                "total_cost": "Total Cost"
            },
            use_container_width=True
        )
    else:
        st.info("No faculty travel data available.")