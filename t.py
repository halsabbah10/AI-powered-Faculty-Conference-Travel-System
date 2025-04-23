import json
import uuid
import docx
import streamlit as st
import mysql.connector
import os
import pandas as pd
import hashlib
from dotenv import load_dotenv
import PyPDF2
from datetime import datetime, timedelta
import requests
from openai import OpenAI
import logging
from bs4 import BeautifulSoup
from functools import lru_cache
import smtplib
from email.mime.text import MIMEText
import plotly.express as px  # type: ignore
import plotly.graph_objects as go  # type: ignore
from logging.handlers import RotatingFileHandler
import traceback
from io import BytesIO
from docx import Document
import re
# Add this connection pooling implementation near the top of your file
import mysql.connector.pooling
from google import genai
from google.genai import types
from pylatex import Document
from datetime import date
import httpx
# Create a connection pool
db_config = {
    "host": "localhost",
    "user": "root",
    "password": "root",
    "database": "con_system",
    "pool_name": "con_system_pool",
    "pool_size": 5,
}
purpose=["Attending a Conference","Presenting at a Conference"]
indexes=["Scopus", "IEEE", "Web of Science", "PubMed", "MEDLINE", "ACM", "None"]
location = {
    "USA": ["New York", "Los Angeles", "Chicago"],
    "Canada": ["Toronto", "Vancouver", "Montreal"],
    "Germany": ["Berlin", "Munich", "Hamburg"],
}
# Initialize the connection pool
try:
    connection_pool = mysql.connector.pooling.MySQLConnectionPool(**db_config)
    logging.info("Database connection pool created successfully")
except Exception as e:
    logging.error(f"Error creating connection pool: {str(e)}")
    connection_pool = None


def get_db_connection():
    """Establish and return a database connection"""
    try:
        return mysql.connector.connect(
            user=os.getenv("DB_USER"),
            password=os.getenv("DB_PASS"),
            host="localhost",
            database=os.getenv("DB_NAME"),
            auth_plugin="mysql_native_password",
        )
    except mysql.connector.Error as e:
        # Handle common connection errors with specific messages
        if e.errno == 2003:  # Can't connect to MySQL server
            st.error("Cannot connect to the database server. Ensure MySQL is running.")
        elif e.errno == 1045:  # Access denied
            st.error("Database access denied. Check your username and password.")
        elif e.errno == 1049:  # Unknown database
            st.error(f"Database '{os.getenv('DB_NAME')}' does not exist.")
        else:
            st.error(f"Database connection error: {str(e)}")

        st.stop()  # Stop execution to prevent further errors



# Must be the first Streamlit command
st.set_page_config(
    page_title="Conference Travel System",
    page_icon="✈️",
    layout="wide",
    initial_sidebar_state="expanded",
)


# Consolidated CSS - place at beginning of the file after imports
def load_css():
    """Load consolidated CSS styles with improved organization"""
    # Base colors and variables
    st.markdown(
        """
        <style>
        /* Color Palette */
        :root {
            --primary: #3C313D;          /* Deep plum */
            --primary-light: #55535B;    /* Lighter plum */
            --primary-dark: #18141D;     /* Dark plum */
            --secondary: #9D9BA2;        /* Subtle grey */
            --background: #EBEAEC;       /* Light grey background */
            --surface: #FFFFFF;          /* White */
            --text-primary: #18141D;     /* Dark text */
            --text-secondary: #55535B;   /* Secondary text */
            --text-on-primary: #FFFFFF;  /* White text on dark */
            --text-on-surface: #18141D;  /* Dark text on light */
            --border: #9D9BA2;          /* Border color */
            --success: #556B2F;         /* Subtle olive green */
            --warning: #8B7355;         /* Muted brown */
            --error: #8B4513;           /* Saddle brown */
        }
        
        /* Typography */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        
        html, body, [class*="css"] {
            font-family: 'Inter', sans-serif;
            color: var(--text-primary);
            background-color: var(--background);
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Component styles
    st.markdown(
        """
        <style>
        /* Main Header */
        .main-header {
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
            /*color: white;*/
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        }
        
        .main-header h1 {
            color: var(--text-on-primary);
            font-weight: 700;
            margin-bottom: 0.5rem;
            font-size: 2.2rem;
        }
        
        .main-header p {
            opacity: 0.9;
            font-size: 1.1rem;
        }
        
        /* Cards and Containers */
        .custom-card, .stat-card {
            background: var(--surface);
            padding: 1.5rem;
            border-radius: 10px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.05);
            margin-bottom: 1.5rem;
            border: 1px solid rgba(0,0,0,0.05);
            transition: transform 0.2s ease, box-shadow 0.2s ease;
        }
        
        .custom-card:hover, .stat-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 24px rgba(0,0,0,0.08);
        }
        
        /* Stats Container */
        .stats-container {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1.5rem;
            margin: 1.5rem 0;
        }
        
        .stat-card {
            text-align: center;
        }
        
        .stat-card h3 {
            color: var(--primary);
            font-size: 1.8rem;
            font-weight: 700;
            margin-bottom: 0.5rem;
        }
        
        .stat-card p {
            color: var(--text-secondary);
            font-size: 1rem;
            font-weight: 500;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # UI Elements and Components
    st.markdown(
        """
        <style>
        /* Buttons */
        div.stButton > button {
            width: 100%;
            padding: 0.75rem 1.5rem;
            border-radius: 8px;
            background: var(--primary);
            color: var(--text-on-primary);
            font-weight: 600;
            border: none;
            transition: all 0.3s ease;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            font-size: 0.9rem;
        }
        
        div.stButton > button:hover {
            background: var(--primary-dark);
            box-shadow: 0 4px 12px rgba(0,0,0,0.2);
            transform: translateY(-2px);
        }
        
        /* Sidebar */
        section[data-testid="stSidebar"] {
            background-color: var(--primary-dark);
            padding: 2rem 0;
        }
        
        section[data-testid="stSidebar"] .stButton > button {
            background-color: transparent;
            color: var(--text-on-primary) !important;
            text-align: left;
            padding: 1rem 2rem;
            border-radius: 0;
            border-left: 4px solid transparent;
            text-transform: none;
            letter-spacing: normal;
            width: 100%;
        }
        
        section[data-testid="stSidebar"] .stButton > button:hover,
        section[data-testid="stSidebar"] .stButton > button:focus,
        section[data-testid="stSidebar"] .stButton > button:active {
            background-color: rgba(255,255,255,0.1);
            border-left: 4px solid var(--primary-light);
            transform: none;
        }
        
        /* Form Elements */
        div.stTextInput > div > div > input,
        div.stTextArea > div > div > textarea,
        div.stSelectbox > div > div > select {
            border-radius: 8px;
            border: 2px solid var(--border);
            padding: 0.75rem 1rem;
            font-size: 1rem;
            transition: all 0.3s ease;
            background: var(--surface);
            /*color: white !important;*/
        }
        
        /* Focus States */
        div.stTextInput > div > div > input:focus,
        div.stTextArea > div > div > textarea:focus {
            border-color: var(--primary);
            box-shadow: 0 0 0 2px rgba(60,49,61,0.1);
        }
        
        /* Charts */
        .js-plotly-plot {
            border-radius: 8px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.05);
            padding: 1rem;
            background: var (--surface);
        }
        
        /* Fix chart text visibility */
        .js-plotly-plot .plotly .main-svg text {
            fill: var(--text-primary) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Special element styling
    st.markdown(
        """
        <style>
        /* Tab styling */
        .stTabs [data-baseweb="tab-list"] {
            gap: 2rem;
            border-bottom: 2px solid var(--border);
        }
        
        .stTabs [data-baseweb="tab"] {
            padding: 1rem 2rem;
            color: var(--text-on-primary) !important;
            font-weight: 500;
        }
        
        .stTabs [data-baseweb="tab"][aria-selected="true"] {
            color: var(--text-on-primary) !important;
            border-bottom: 4px solid var(--primary);
            font-weight: 600;
        }
        
        /* Status Tags */
        .status-tag {
            display: inline-block;
            padding: 0.25rem 0.75rem;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 500;
        }
        
        .status-approved { background: rgba(85,107,47,0.1); color: var(--success); }
        .status-pending { background: rgba(139,115,85,0.1); color: var(--warning); }
        .status-rejected { background: rgba(139,69,19,0.1); color: var(--error); }
        
        /* Ensure dark backgrounds have light text */
        section[data-testid="stSidebar"], 
        .main-header, 
        .streamlit-expander > div:first-child {
            color: var(--text-on-primary) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Add this CSS snippet after your other CSS definitions (around line 645)
    st.markdown(
        """
        <style>
        /* Override text input colors - make text white */
        div.stTextInput > div > div > input,
        div.stTextArea > div > textarea,
        div.stSelectbox > div > div > select {
            /*color: white !important; */
            /*caret-color: white !important;*/
        }

        /* Make sure placeholder text is also visible */
        div.stTextInput > div > div > input::placeholder,
        div.stTextArea > div > div > textarea::placeholder {
            color: rgba(255, 255, 255, 0.7) !important;
        }
        
        /* Background color for inputs on dark themes */
        div.stTextInput > div > div > input,
        div.stTextArea > div > div > textarea,
        div.stSelectbox > div > div > select {
            background-color: rgba(255, 255, 255, 0.1) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# Call this function early in your app
load_css()


def setup_logging(log_level=logging.INFO):
    """Set up comprehensive logging system"""
    # Create logs directory if it doesn't exist
    if not os.path.exists("logs"):
        os.makedirs("logs")

    # Define log file with date
    log_filename = f"logs/app_{datetime.now().strftime('%Y-%m-%d')}.log"

    # Configure logging with both file and console output
    handlers = [
        # Rotating file handler (10MB per file, keep 10 files)
        RotatingFileHandler(log_filename, maxBytes=10_485_760, backupCount=10),
        # Console handler
        logging.StreamHandler(),
    ]

    # Define detailed log format
    log_format = "%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s"
    date_format = "%Y-%m-%d %H:%M:%S"

    # Configure logging
    logging.basicConfig(
        level=log_level, format=log_format, datefmt=date_format, handlers=handlers
    )

    # Log startup information
    logging.info(f"Application started. Log level: {logging.getLevelName(log_level)}")


def log_error(error, context=None):
    """Log error with detailed information and context"""
    tb_str = traceback.format_exception(type(error), error, error.__traceback__)
    error_details = {
        "error_type": error.__class__.__name__,
        "error_message": str(error),
        "traceback": "".join(tb_str),
        "context": context or {},
    }

    logging.error(
        f"Error: {error_details['error_type']}: {error_details['error_message']}"
    )
    logging.error(f"Context: {error_details['context']}")
    logging.debug(f"Traceback: {error_details['traceback']}")

    return error_details


# Call setup_logging() at the beginning of your application
if __name__ == "__main__":
    setup_logging()


def sanitize_input(text):
    """Sanitize input text to prevent injection attacks."""
    return text.strip()


def validate_url(url):
    """Validate URL by fully mimicking a browser to avoid disconnects."""
    try:
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            ),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Connection": "keep-alive",
        }

        with httpx.Client(timeout=20, follow_redirects=True) as client:
            response = client.get(url, headers=headers)
            return response.status_code == 200
    except httpx.ReadError as e:
        print(f"URL check failed: Server disconnected without response. ({e})")
        return False
    except httpx.RequestError as e:
        print(f"URL check failed: {e}")
        return False

@lru_cache(maxsize=128)
def extract_conference_info(url):
    """Extract text content from a conference website."""
    try:
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.content, "html.parser")
        main_content = soup.find_all(["main", "article", "div.content"])
        return " ".join(
            [elem.get_text(separator=" ", strip=True) for elem in main_content]
        )
    except Exception as e:
        logging.error(f"Error extracting content: {e}")
        return ""


def send_email_notification(recipient, subject, message):
    """Send an email notification via SMTP."""
    smtp_server = os.getenv("SMTP_SERVER")
    smtp_port = int(os.getenv("SMTP_PORT", 587))
    smtp_username = os.getenv("SMTP_USERNAME")
    smtp_password = os.getenv("SMTP_PASSWORD")
    sender = os.getenv("EMAIL_SENDER")
    try:
        msg = MIMEText(message)
        msg["Subject"] = subject
        msg["From"] = sender
        msg["To"] = recipient
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_username, smtp_password)
            server.sendmail(sender, [recipient], msg.as_string())
        logging.info(f"Email sent to {recipient}: {subject}")
    except Exception as e:
        logging.error(f"Email sending failed: {e}")


# Load environment variables
load_dotenv()
client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),  # This is the default and can be omitted
)


# Acceptance Letter Summary Function
def generate_conference_summary(acceptance_text, conference_name):
    """Generate a detailed summary of the conference acceptance letter."""
    if not acceptance_text:
        return "Invalid: No acceptance letter provided."

    try:
        prompt = f"""
        Analyze this conference acceptance letter for '{conference_name}' and provide a structured summary:
        1. Verification: Confirm if this is a valid acceptance letter
        2. Conference Details: Name, dates, location
        3. Presentation Type: Oral, poster, etc.
        4. Key Requirements: Registration deadlines, presentation requirements
        5. Any red flags or concerns

        Letter text:
        {acceptance_text}
        """

        response = client.chat.completions.create(
            model="gpt-4",  # Using GPT-4 for better analysis
            messages=[
                {
                    "role": "system",
                    "content": "You are a conference verification specialist who analyzes acceptance letters.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=500,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logging.error(f"Error generating conference summary: {e}")
        return f"Error during verification: {str(e)}"


def generate_ai_notes(conference_name, purpose, index_type, destination_country, city):
    """Generate AI notes about the conference request."""
    try:
        prompt = f"""
        Analyze this conference travel request and provide insights:
        Conference: {conference_name}
        Purpose: {purpose}
        Index Type: {index_type}
        Location: {city}, {destination_country}

        Please provide:
        1. Travel Impact: Value of attending this conference
        2. Location Analysis: Benefits/challenges of the destination
        3. Index Quality: Assessment of the conference indexing
        4. Recommendations: Specific suggestions for the traveler
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system",
                    "content": "You are a travel and academic advisor who provides insights on conference trips.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=400,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logging.error(f"Error generating AI notes: {e}")
        return f"Error generating insights: {str(e)}"

def extract_text_from_file(file):
    """Extract text from uploaded PDF or Word documents"""
    try:
        if file.type == "application/pdf":
            # Handle PDF files
            pdf_reader = PyPDF2.PdfReader(BytesIO(file.getvalue()))
            return "\n".join([page.extract_text() for page in pdf_reader.pages])
        
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Handle Word documents
            doc = docx.Document(BytesIO(file.getvalue()))
            return "\n".join([para.text for para in doc.paragraphs])
        
        else:
            raise ValueError("Unsupported file format")
            
    except Exception as e:
        print(f"Error extracting text from file: {str(e)}")
        raise

@lru_cache(maxsize=128)
def extract_index_from_url(url, selected_index):
    """Simplified website verification checking only for conference mentions"""
    try:
        response = requests.get(url, timeout=15)
        soup = BeautifulSoup(response.content, 'html.parser')
        text = ' '.join([p.get_text() for p in soup.find_all(['p', 'h1', 'h2', 'h3', 'li'])])[:2000]

        # Case-insensitive pattern matching
        pattern = re.compile(fr'\b{re.escape(selected_index)}\b', re.IGNORECASE)
        
        return bool(pattern.search(text)), [selected_index]

    except Exception as e:
        print(f"Simplified check error: {str(e)}")
        return False, []




def validate_with_gpt(conference_name, research_paper_text, acceptance_letter_text, selected_index, conference_url=None):
    """Validate conference submission documents with automated index verification"""
    # URL-based index check
    url_index_match, found_indexes = (False, [])
    if conference_url:
        url_index_match, found_indexes = extract_index_from_url(conference_url, selected_index)

    # Check for organizational naming convention matches
    org_in_name = any(org in conference_name.upper() for org in ["ACM", "IEEE"])
    org_match = selected_index.upper() in conference_name.upper()

    prompt = f"""
    Validate these conference submission elements:
    1. Conference: {conference_name} (Organization in name: {'Yes' if org_in_name else 'No'})
    2. Selected Index: {selected_index}
    3. Paper Content: {research_paper_text[:700]}
    4. Acceptance Letter: {acceptance_letter_text[:700]}
    5. Website Index Mentions: {', '.join(found_indexes) or 'None found'}
    
    Perform these checks:
    - Verify acceptance letter authenticity through organizational details
    - Confirm research paper aligns with conference's stated themes
    - Check ANY mention of {selected_index} in website/content
    - Consider organizational naming conventions (e.g., 'ACM' in name)
    - Allow alternate verification methods for organizational conferences
    
    Return JSON with structure:
    {{
        "valid": boolean,
        "message": "Summary including organizational affiliation consideration",
        "issues": ["critical issues"] or [],
        "index_verified": boolean|null
    }}
    Respond ONLY with valid JSON"""

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": """You are an academic validation assistant that considers:
                - Organizational naming conventions
                - Multiple verification methods
                - Conference reputation indicators"""},
                {"role": "user", "content": prompt}
            ],
            temperature=0
        )
        
        validation_result = response.choices[0].message.content

        
        try:
            result = json.loads(validation_result)
            required_keys = ["valid", "message", "issues", "index_verified"]
            if not all(key in result for key in required_keys):
                return {
                    "valid": False,
                    "message": "Invalid validation structure",
                    "issues": ["Missing required fields"],
                    "index_verified": None
                }

            # Enhanced Index Verification Logic
            if org_match:
                # Conference name contains selected index organization
                if not result["index_verified"]:
                    result["issues"] = [f"Recommended {selected_index} verification missing"]
                    result["valid"] = True  # Override to valid with warning
                    result["index_verified"] = True  # Consider verified by naming
            elif selected_index in ["ACM", "IEEE"]: 
                if not result["index_verified"]:
                    result["issues"].append(f"Recommended {selected_index} verification missing")
                    result["valid"] = True  # Warn instead of reject

            return result
            
        except json.JSONDecodeError as e:
            return {
                "valid": False,
                "message": f"JSON parsing error: {str(e)}",
                "issues": ["Invalid response format"],
                "index_verified": None
            }
            
    except Exception as e:
        logging.error(f"Validation error: {str(e)}")
        return {
            "valid": False,
            "message": "Validation system error",
            "issues": ["API failure"],
            "index_verified": None
        }



# Reaserach paper summary function
def generate_research_summary(text, conference_name):
    """Generate a detailed summary of the research paper."""
    if not text:
        return "Invalid: No research paper provided."

    try:
        prompt = f"""
        Analyze this research paper for the conference '{conference_name}' and provide a structured summary:
        1. Research Topic: Main focus of the paper
        2. Methodology: Key research methods used
        3. Findings: Principal results and conclusions
        4. Relevance: Connection to conference theme
        5. Quality Assessment: Academic rigor and contribution

        Paper text:
        {text}
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system",
                    "content": "You are a research paper analyst who evaluates academic papers.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=500,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logging.error(f"Error generating research summary: {e}")
        return f"Error analyzing research paper: {str(e)}"


# Initialize session state for user login and page navigation
if "logged_in_user" not in st.session_state:
    st.session_state.logged_in_user = None  # Initialize logged-in user as None
if "page" not in st.session_state:
    st.session_state.page = "login"  # Set initial page to login
if "user_role" not in st.session_state:
    st.session_state.user_role = None  # Track user role for permissions
if "refresh_budget" not in st.session_state:
    st.session_state.refresh_budget = (
        False  # Dummy variable for refreshing budget table
    )


class DatabaseManager:
    @staticmethod
    def execute_query(query, params=None, fetch=True, commit=False):
        """Execute database query with proper connection management"""
        connection = None
        cursor = None
        result = None
        try:
            connection = get_db_connection()
            cursor = connection.cursor(dictionary=True)
            cursor.execute(query, params or ())

            if fetch:
                result = cursor.fetchall()
            if commit:
                connection.commit()
                result = cursor.lastrowid

            return result
        except mysql.connector.Error as e:
            logging.error(f"Database error: {str(e)}")
            if commit:
                connection.rollback()
            raise
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()


# Function to check ID and password
def check_credentials(user_id, password):
    connection = get_db_connection()
    cursor = connection.cursor(dictionary=True)
    hashed_password = hashlib.sha256(password.encode()).hexdigest()
    cursor.execute(
        "SELECT * FROM faculty WHERE user_id = %s AND password = %s",
        (user_id, hashed_password),
    )
    result = cursor.fetchone()
    cursor.close()
    connection.close()
    if result:
        st.session_state.user_role = result["role"]  # Set user role in session state
        st.session_state.user_name = result[
            "name"
        ]  # Store user's name in session state
        return True
    return False


def login_callback():
    """Handle user login"""
    user_id = st.session_state.get("user_id_input")
    password = st.session_state.get("password_input")

    if user_id and password and check_credentials(user_id, password):
        # Clear previous session data
        for key in list(st.session_state.keys()):
            if key not in ["page", "logged_in_user", "user_role", "user_name"]:
                del st.session_state[key]

        # Set new session data
        st.session_state.logged_in_user = user_id
        st.session_state.page = "main"
        st.session_state.submitted = False
        st.session_state.login_attempted = True
        st.session_state.show_success = True

    else:
        st.error("Invalid ID or password.")
        st.session_state.login_attempted = False


def is_within_restricted_dates(start_date, end_date):
    """
    Checks if selected travel dates overlap with restricted periods.
    Provides a message if users select restricted dates.
    """
    connection = get_db_connection()
    cursor = connection.cursor(dictionary=True)

    # Query to fetch all restricted periods
    cursor.execute("SELECT start_date, end_date, description FROM restricteddates")
    restricted_periods = cursor.fetchall()
    cursor.close()
    connection.close()

    # Iterate through restricted periods to check for conflicts
    conflict_messages = []
    for period in restricted_periods:
        restricted_start = period["start_date"]
        restricted_end = period["end_date"]

        # If selected dates overlap with restricted dates, append conflict info
        if start_date <= restricted_end and end_date >= restricted_start:
            conflict_messages.append(
                f"{period['description']} ({restricted_start} to {restricted_end})"
            )

    # If conflicts exist, return them; otherwise, return None
    if conflict_messages:
        return False, "Travel dates overlap with restricted periods: " + ", ".join(
            conflict_messages
        )
    return True, None


# Display date conflict warning without disabling form


def has_traveled_this_year(user_id):
    connection = get_db_connection()
    cursor = connection.cursor(dictionary=True)

    # Current year
    current_year = datetime.now().year

    # Check for any approved requests for the current user in the current year
    cursor.execute(
        """
        SELECT COUNT(*) as travel_count FROM requests 
        WHERE faculty_user_id = %s AND status = 'approved'
        AND YEAR(date_from) = %s
    """,
        (user_id, current_year),
    )

    result = cursor.fetchone()
    cursor.close()
    connection.close()

    # Return True if there's at least one approved travel request in the current year
    return result["travel_count"] > 0


# Updated login page with ID and password
def login_page():
    st.header("Login")
    
    # Use session state to store input fields for callback access
    st.text_input("Enter your ID", key="user_id_input")
    st.text_input("Enter your password", type="password", key="password_input", on_change=login_callback)
    
    # Login button with callback to handle login logic
    st.button("Login", on_click=login_callback)
# Function to get request counts and details for the professor's home page
def fetch_user_request_counts(user_id):
    connection = get_db_connection()
    cursor = connection.cursor(dictionary=True)

    # Get request counts by status
    cursor.execute(
        """
        SELECT 
            COUNT(CASE WHEN status = 'pending' THEN 1 END) AS pending_count,
            COUNT(CASE WHEN status = 'approved' THEN 1 ELSE 0 END) AS approved_count,
            COUNT(*) AS total_count
        FROM requests
        WHERE faculty_user_id = %s
    """,
        (user_id,),
    )
    counts = cursor.fetchone()

    # Get detailed request data for display, using conference_name from Requests
    cursor.execute(
        """
        SELECT request_id, conference_name, status, submission_date 
        FROM requests
        WHERE faculty_user_id = %s
    """,
        (user_id,),
    )
    request_details = cursor.fetchall()

    cursor.close()
    connection.close()
    return counts, request_details


# Fetch current available budget, converted to float
@st.cache_data(ttl=300)
def fetch_available_budget():
    """Fetch available budget with caching"""
    result = DatabaseManager.execute_query("SELECT funds_available FROM budget LIMIT 1")
    return float(result[0]["funds_available"]) if result else 0.0


# Fetch budget history
def fetch_budget_history():
    connection = get_db_connection()
    cursor = connection.cursor(dictionary=True)
    cursor.execute(
        "SELECT adjustment_type, amount, description, update_timestamp FROM budgethistory ORDER BY update_timestamp DESC"
    )
    history_data = cursor.fetchall()
    cursor.close()
    connection.close()
    return history_data


# Function to set or adjust the available budget
def set_budget(new_budget):
    connection = get_db_connection()
    cursor = connection.cursor()

    # Update or insert the budget and add an adjustment entry to the history
    cursor.execute(
        """
        INSERT INTO budget (budget_id, funds_available) VALUES (1, %s)
        ON DUPLICATE KEY UPDATE funds_available = %s
    """,
        (new_budget, new_budget),
    )

    cursor.execute(
        """
        INSERT INTO budgethistory (adjustment_type, amount, description)
        VALUES ('Adjustment', %s, 'Budget set or adjusted by accountant')
    """,
        (new_budget,),
    )

    connection.commit()
    cursor.close()
    connection.close()


# Function to deduct travel expenses from the available budget
def deduct_from_budget(expense, description):
    connection = get_db_connection()
    cursor = connection.cursor()

    # Deduct from budget and log the deduction in BudgetHistory
    cursor.execute(
        "UPDATE budget SET funds_available = funds_available - %s WHERE budget_id = 1",
        (expense,),
    )
    cursor.execute(
        """
        INSERT INTO budgethistory (adjustment_type, amount, description)
        VALUES ('Deduction', %s, %s)
    """,
        (expense, description),
    )

    connection.commit()
    cursor.close()
    connection.close()


# Function to approve request and deduct the cost from budget
def approve_request_and_deduct_budget(request_id, total_cost):
    """Approve request and deduct budget with improved error handling"""
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        cursor.execute("START TRANSACTION")

        # Update request status
        cursor.execute(
            "UPDATE requests SET status = 'approved' WHERE request_id = %s", (request_id,)
        )

        # Check if budget is sufficient
        cursor.execute(
            "SELECT funds_available FROM budget WHERE budget_id = 1 FOR UPDATE"
        )
        available_budget = cursor.fetchone()[0]

        if available_budget < total_cost:
            cursor.execute("ROLLBACK")
            logging.warning(
                f"Insufficient budget for request {request_id}: {total_cost} > {available_budget}"
            )
            return False, "Insufficient budget"

        # Deduct from budget
        cursor.execute(
            "UPDATE budget SET funds_available = funds_available - %s WHERE budget_id = 1",
            (total_cost,)  # Ensure total_cost is passed as a tuple
        )

        # Add history record
        cursor.execute(
            """INSERT INTO budgethistory (adjustment_type, amount, description)
               VALUES ('Deduction', %s, %s)""",
            (total_cost, f"Deduction for approved request ID {request_id}")
        )

        # Commit transaction
        cursor.execute("COMMIT")

        return True, "Request approved successfully"

    except Exception as e:
        if "connection" in locals() and connection:
            cursor.execute("ROLLBACK")

        # Log detailed error information if needed
        logging.error(f"Error occurred while processing request {request_id}: {str(e)}")

        return False, f"Error: {str(e)}"

    finally:
        if "cursor" in locals() and cursor:
            cursor.close()
        if "connection" in locals() and connection:
            connection.close()


# Replace the insert_travel_request function
def insert_travel_request(data, acceptance_letter, conference_url, research_paper):
    """Insert travel request with proper connection management"""
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        # Check for an existing request with the same details for this user
        user_id, conference_name, date_from, date_to = (
            data[0],
            data[1],
            data[7],
            data[8],
        )
        cursor.execute(
            """
            SELECT * FROM requests 
            WHERE faculty_user_id = %s AND conference_name = %s AND date_from = %s AND date_to = %s
            """,
            (user_id, conference_name, date_from, date_to),
        )

        existing_request = cursor.fetchone()
        if existing_request:
            cursor.close()
            connection.close()
            return "duplicate"  # Indicate duplicate found

        # Process acceptance letter
        conference_text = ""
        if acceptance_letter is not None:
            try:
                # Reset file pointer to beginning
                acceptance_letter.seek(0)
                reader = PyPDF2.PdfReader(acceptance_letter)
                for page in reader.pages:
                    conference_text += page.extract_text()
            except Exception as e:
                logging.error(f"Error reading acceptance letter PDF: {e}")
                conference_text = ""

        # Generate summary for acceptance letter
        conference_summary = (
            generate_conference_summary(conference_text, data[1])
            if conference_text
            else "Summary unavailable due to missing data"
        )

        # Process research paper
        research_text = ""
        if research_paper is not None:
            try:
                # Reset file pointer to beginning
                research_paper.seek(0)
                reader = PyPDF2.PdfReader(research_paper)
                for page in reader.pages:
                    research_text += page.extract_text()
            except Exception as e:
                logging.error(f"Error reading research paper PDF: {e}")
                research_text = ""

        # Generate summary for the research paper
        research_summary = (
            generate_research_summary(research_text, data[1])
            if research_text
            else "Summary unavailable due to missing data"
        )

        # Extract fields for AI notes
        conference_name = data[1]
        purpose_of_attending = data[2]
        index_type = data[3]
        destination_country = data[4]
        city = data[5]

        # Generate AI notes with correct fields
        notes_summary = generate_ai_notes(
            conference_name, purpose_of_attending, index_type, destination_country, city
        )

        # Ensure data tuple has correct length
        if len(data) != 13:
            raise ValueError("The data tuple must include 13 elements.")

        # Insert request data along with summaries and AI notes
        cursor.execute(
            """
            INSERT INTO requests (faculty_user_id, conference_name, status, purpose_of_attending, index_type,
                                  destination, city, date_from, date_to, per_diem, registration_fee, visa_fee,
                                  conference_url, url_summary, submission_date, conference_summary, research_summary, notes_summary)
            VALUES (%s, %s, 'pending', %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), %s, %s, %s)
            """,
            data + (conference_summary, research_summary, notes_summary),
        )

        # Fetch the last inserted ID
        request_id = cursor.lastrowid
        connection.commit()

        return request_id

    except Exception as e:
        logging.error(f"Error inserting travel request: {str(e)}")
        if "connection" in locals() and connection:
            connection.rollback()
        return None

    finally:
        if "cursor" in locals() and cursor:
            cursor.close()
        if "connection" in locals() and connection:
            connection.close()


# Insert uploaded files into UploadedFiles table
def insert_uploaded_files(request_id, acceptance_letter, research_paper):
    connection = get_db_connection()
    cursor = connection.cursor()

    try:
        if acceptance_letter is not None:
            acceptance_letter.seek(0)
            file_bytes = acceptance_letter.read()  # Read once
            cursor.execute(
                """
                INSERT INTO uploadedfiles (request_id, file_data, file_type, file_path)
                VALUES (%s, %s, %s, %s)
            """,
                (request_id, file_bytes, "acceptance_letter", ""),
            )

        if research_paper is not None:
            research_paper.seek(0)
            file_bytes = research_paper.read()  # Read once
            cursor.execute(
                """
                INSERT INTO uploadedfiles (request_id, file_data, file_type, file_path)
                VALUES (%s, %s, %s, %s)
            """,
                (request_id, file_bytes, "research_paper", ""),
            )

        connection.commit()

    except Exception as e:
        print(f"Error inserting uploaded files: {e}")

    finally:
        cursor.close()
        connection.close()


# Fetch pending travel requests for approval
def fetch_pending_requests():
    connection = get_db_connection()
    cursor = connection.cursor(dictionary=True)
    cursor.execute("""
        SELECT requests.request_id, faculty.name AS requester, 
               requests.conference_name, requests.purpose_of_attending, 
               requests.conference_url, requests.url_summary, requests.destination, 
               requests.city, requests.date_from, requests.date_to, 
               requests.per_diem, requests.registration_fee, requests.visa_fee, 
               requests.conference_summary, requests.research_summary, requests.notes_summary,
               acceptance.file_data AS acceptance_letter, 
               research.file_data AS research_paper
        FROM requests
        JOIN faculty ON requests.faculty_user_id = faculty.user_id
        LEFT JOIN uploadedfiles AS acceptance 
            ON acceptance.request_id = requests.request_id AND acceptance.file_type = 'acceptance_letter'
        LEFT JOIN uploadedfiles AS research 
            ON research.request_id = requests.request_id AND research.file_type = 'research_paper'
        WHERE requests.status = 'pending'
    """)
    pending_requests = cursor.fetchall()
    cursor.close()
    connection.close()
    return pending_requests


def process_url(conference_url):
    """Process conference URL to extract a brief summary."""
    if not conference_url:
        return None
    info = extract_conference_info(conference_url)
    if info:
        return info[:300] + "..." if len(info) > 300 else info
    return None


def record_audit_log(user_id, action, details):
    """Record an audit log entry."""
    log_entry = f"{datetime.now().isoformat()} - User: {user_id} - Action: {action} - Details: {details}\n"
    with open("audit.log", "a") as log_file:
        log_file.write(log_entry)
    logging.info(log_entry)


# Add this new function for dashboard data
@st.cache_data(ttl=300)
def fetch_dashboard_data():
    """Fetch comprehensive dashboard data with improved error handling"""
    try:
        # Monthly stats with better detail
        monthly_stats = DatabaseManager.execute_query("""
            SELECT 
                DATE_FORMAT(submission_date, '%Y-%m') as month,
                COUNT(*) as total_requests,
                SUM(CASE WHEN status = 'approved' THEN 1 ELSE 0 END) as approved_requests,
                SUM(CASE WHEN status = 'pending' THEN 1 ELSE 0 END) as pending_requests,
                SUM(per_diem + registration_fee + visa_fee) as total_cost
            FROM requests
            GROUP BY DATE_FORMAT(submission_date, '%Y-%m')
            ORDER BY month
        """)

        # Status distribution
        status_stats = DatabaseManager.execute_query("""
            SELECT status, COUNT(*) as count
            FROM requests
            GROUP BY status
        """)

        # Top destinations with costs
        destination_stats = DatabaseManager.execute_query("""
            SELECT 
                destination,
                COUNT(*) as visit_count,
                SUM(per_diem + registration_fee + visa_fee) as total_cost
            FROM requests
            GROUP BY destination
            ORDER BY visit_count DESC
            LIMIT 10
        """)

        # Faculty travel frequency
        faculty_stats = DatabaseManager.execute_query("""
            SELECT 
                f.name,
                COUNT(r.request_id) as travel_count,
                SUM(r.per_diem + r.registration_fee + r.visa_fee) as total_spent
            FROM faculty f
            LEFT JOIN requests r ON f.user_id = r.faculty_user_id
            WHERE r.status = 'approved'
            GROUP BY f.name
            ORDER BY travel_count DESC
        """)

        return {
            "monthly_stats": monthly_stats or [],
            "status_stats": status_stats or [],
            "destination_stats": destination_stats or [],
            "faculty_stats": faculty_stats or [],
        }
    except Exception as e:
        logging.error(f"Error fetching dashboard data: {str(e)}")
        st.error("Failed to load dashboard data. Please try again later.")
        return {
            "monthly_stats": [],
            "status_stats": [],
            "destination_stats": [],
            "faculty_stats": [],
        }


def create_dashboard_charts(data):
    """Create enhanced dashboard visualizations"""

    # Common chart config
    chart_config = create_chart_config()

    # Monthly Trends - Fix data type issue
    monthly_df = pd.DataFrame(data["monthly_stats"])
    if not monthly_df.empty:
        # Convert month to datetime for proper ordering
        monthly_df["month"] = pd.to_datetime(monthly_df["month"])

        # Create separate traces for each metric
        fig_trends = go.Figure()

        # Add each metric as a separate line
        metrics = ["total_requests", "approved_requests", "pending_requests"]
        colors = ["#3C313D", "#556B2F", "#8B7355"]

        for metric, color in zip(metrics, colors):
            fig_trends.add_trace(
                go.Scatter(
                    x=monthly_df["month"],
                    y=monthly_df[metric],
                    name=metric.replace("_", " ").title(),
                    line=dict(color=color, width=2),
                )
            )

        fig_trends.update_layout(
            **chart_config,
            title="Monthly Request Trends",
            xaxis_title="Month",
            yaxis_title="Number of Requests",
        )
        st.plotly_chart(fig_trends, use_container_width=True)

    # Create two-column layout
    col1, col2 = st.columns(2)

    with col1:
        # Status Distribution
        status_df = pd.DataFrame(data["status_stats"])
        if not status_df.empty:
            fig_status = px.pie(
                status_df,
                values="count",
                names="status",
                title="Request Status Distribution",
                color_discrete_sequence=["#3C313D", "#556B2F", "#8B7355"],
            )
            fig_status.update_layout(**chart_config)
            st.plotly_chart(fig_status, use_container_width=True)

    with col2:
        # Destination Analysis
        dest_df = pd.DataFrame(data["destination_stats"])
        if not dest_df.empty:
            fig_dest = px.bar(
                dest_df,
                x="destination",
                y="visit_count",
                title="Top Destinations",
                color="total_cost",
                color_continuous_scale="Purples",
            )
            fig_dest.update_layout(**chart_config)
            fig_dest.update_xaxes(tickangle=45)
            st.plotly_chart(fig_dest, use_container_width=True)

    # Faculty Travel Analysis
    faculty_df = pd.DataFrame(data["faculty_stats"])
    if not faculty_df.empty:
        # Separate the metrics into different visualizations
        fig_faculty = go.Figure()

        # Add travel count bars
        fig_faculty.add_trace(
            go.Bar(
                x=faculty_df["name"],
                y=faculty_df["travel_count"],
                name="Travel Count",
                marker_color="#3C313D",
            )
        )

        # Add total spent bars
        fig_faculty.add_trace(
            go.Bar(
                x=faculty_df["name"],
                y=faculty_df["total_spent"],
                name="Total Spent (AED)",
                marker_color="#556B2F",
                yaxis="y2",
            )
        )

        fig_faculty.update_layout(
            **chart_config,
            title="Faculty Travel Analysis",
            xaxis_title="Faculty Member",
            yaxis_title="Travel Count",
            yaxis2=dict(title="Total Spent (AED)", overlaying="y", side="right"),
            barmode="group",
        )

        fig_faculty.update_xaxes(tickangle=45)
        st.plotly_chart(fig_faculty, use_container_width=True)


def calculate_metrics(data: pd.DataFrame) -> dict:
    """
    Calculate key metrics from the provided DataFrame.

    Args:
        data (pd.DataFrame): Input DataFrame containing conference travel data

    Returns:
        dict: Dictionary containing calculated metrics
    """
    try:
        metrics = {
            "total_conferences": len(data),
            "total_budget": data["budget"].sum() if "budget" in data.columns else 0,
            "avg_cost": data["budget"].mean() if "budget" in data.columns else 0,
            "pending_approvals": len(data[data["status"] == "pending"])
            if "status" in data.columns
            else 0,
            "approved_requests": len(data[data["status"] == "approved"])
            if "status" in data.columns
            else 0,
        }
        return metrics

    except Exception as e:
        logging.error(f"Error calculating metrics: {str(e)}")
        st.error("Failed to calculate metrics. Please check the data format.")
        return {}


def display_pending_requests():
    """Display only pending requests without analytics"""
    pending_requests = fetch_pending_requests()

    if not pending_requests:
        st.write("No pending travel requests for approval.")
    else:
        for request in pending_requests:
            with st.expander(
                f"Request ID: {request['request_id']} - {request['requester']}"
            ):
                # Display request details
                st.markdown(f"**Applicant:** {request['requester']}")
                st.markdown(f"**Conference:** {request['conference_name']}")
                st.markdown(f"**Purpose:** {request['purpose_of_attending']}")
                st.markdown(
                    f"**Location:** {request['destination']}, {request['city']}"
                )
                st.markdown(f"**Dates:** {request['date_from']} - {request['date_to']}")

                # Cost breakdown
                total_cost = (
                    request["per_diem"]
                    + request["registration_fee"]
                    + request["visa_fee"]
                )
                st.markdown("#### Cost Breakdown")
                st.markdown(f"- Per Diem: AED {request['per_diem']:,.2f}")
                st.markdown(f"- Registration Fee: AED {request['registration_fee']:,.2f}")
                st.markdown(f"- Visa Fee: AED {request['visa_fee']:,.2f}")
                st.markdown(f"**Total Cost:** AED {total_cost:,.2f}")

                # AI summaries
                st.markdown("**AI-Generated Summary & Notes**")
                st.markdown("**Acceptance Letter Summary:**")
                st.markdown(request.get('conference_summary', 'Summary unavailable'))

                st.markdown("**Research Paper Summary:**")
                st.markdown(request.get('research_summary', 'Research summary unavailable'))

                # Downloadable files
                if request.get('acceptance_letter'):
                    st.download_button(
                        label="Download Acceptance Letter",
                        data=request['acceptance_letter'],
                        file_name="acceptance_letter.pdf",
                        mime="application/pdf",  # ensures the file is handled as a PDF
                        key=f"dl_accept_{request['request_id']}"
                    )

                if request.get('research_paper'):
                    st.download_button(
                        label="Download Research Paper",
                        data=request['research_paper'],
                        file_name="research_paper.pdf",
                        mime="application/pdf",
                        key=f"dl_research_{request['request_id']}"
                    )


                # Approval buttons
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Approve", key=f"approve_{request['request_id']}"):
                        approve_request_and_deduct_budget(
                            request["request_id"], total_cost
                        )
                        st.success("Request approved successfully!")
                        st.rerun()
                with col2:
                    if st.button("Reject", key=f"reject_{request['request_id']}"):
                        reject_request(request["request_id"])
                        st.info("Request rejected.")
                        st.rerun()

                # Developer-only delete option
                if st.session_state.logged_in_user == "3":
                    if st.button(f"Delete Request {request['request_id']}", key=f"delete_{request['request_id']}"):
                        st.session_state['delete_request_id'] = request['request_id']

        # Perform the delete action if triggered
        if 'delete_request_id' in st.session_state:
            delete_request_callback()


def calculate_dashboard_metrics(data):
    """Calculate metrics for dashboard display"""
    try:
        total_requests = sum(item["total_requests"] for item in data["monthly_stats"])
        pending_count = sum(
            item["count"]
            for item in data["status_stats"]
            if item["status"] == "pending"
        )
        approved_count = sum(
            item["count"]
            for item in data["status_stats"]
            if item["status"] == "approved"
        )
        total_cost = sum(item["total_cost"] for item in data["monthly_stats"])

        return {
            "total_requests": total_requests,
            "pending_count": pending_count,
            "approved_count": approved_count,
            "total_cost": total_cost,
        }
    except Exception as e:
        st.error(f"Error calculating metrics: {str(e)}")
        return {
            "total_requests": 0,
            "pending_count": 0,
            "approved_count": 0,
            "total_cost": 0,
        }


def create_chart_config(is_monthly_trend=False):
    """Create consistent chart styling with conditional formatting"""
    if is_monthly_trend:
        return {
            "template": "plotly_white",
            "height": 400,
            "margin": dict(l=40, r=40, t=60, b=40),
            "font": dict(family="Inter", size=14, color="#18141D"),
            "paper_bgcolor": "white",
            "plot_bgcolor": "white",
            "showlegend": True,
            "legend": dict(
                orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1
            ),
        }
    else:
        return {
            "template": "plotly_white",
            "height": 400,
            "margin": dict(l=40, r=40, t=60, b=40),
            "font": dict(family="Inter", size=14, color="white"),
            "paper_bgcolor": "#3C313D",
            "plot_bgcolor": "#3C313D",
            "showlegend": True,
            "legend": dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1,
                font=dict(color="white"),
            ),
            "xaxis": dict(tickfont=dict(color="white"), title_font=dict(color="white")),
            "yaxis": dict(tickfont=dict(color="white"), title_font=dict(color="white")),
        }


def display_dashboard():
    """Display unified dashboard with proper formatting"""
    st.markdown("## Dashboard Overview")

    # Fetch all data at once
    dashboard_data = fetch_dashboard_data()

    # Calculate and display metrics
    metrics = calculate_dashboard_metrics(dashboard_data)

    # Display metric cards
    st.markdown(
        f"""
        <div class="stats-container">
            <div class="stat-card">
                <h3>{metrics['total_requests']}</h3>
                <p>Total Requests</p>
            </div>
            <div class="stat-card">
                <h3>{metrics['pending_count']}</h3>
                <p>Pending Approval</p>
            </div>
            <div class="stat-card">
                <h3>{metrics['approved_count']}</h3>
                <p>Approved Requests</p>
            </div>
            <div class="stat-card">
                <h3>AED {metrics['total_cost']:,.2f}</h3>
                <p>Total Cost</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Monthly Request Trends")
    monthly_df = pd.DataFrame(dashboard_data["monthly_stats"])
    if not monthly_df.empty:
        monthly_df["month"] = pd.to_datetime(monthly_df["month"])
        fig = go.Figure()

        metrics = ["total_requests", "approved_requests", "pending_requests"]
        colors = ["#3C313D", "#556B2F", "#8B7355"]

        for metric, color in zip(metrics, colors):
            fig.add_trace(
                go.Scatter(
                    x=monthly_df["month"],
                    y=monthly_df[metric],
                    name=metric.replace("_", " ").title(),
                    line=dict(color=color, width=2),
                )
            )

        fig.update_layout(
            **create_chart_config(is_monthly_trend=True),
            title="Monthly Request Trends",
            xaxis_title="Month",
            yaxis_title="Number of Requests",
        )
        st.plotly_chart(fig, use_container_width=True)

    # Two-column layout for distribution charts
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Request Status Distribution")
        create_status_distribution_chart(dashboard_data["status_stats"])

    with col2:
        st.markdown("### Destination Analysis")
        create_destination_chart(dashboard_data["destination_stats"])

    # Faculty Analysis
    st.markdown("### Faculty Travel Analysis")
    create_faculty_analysis_chart(dashboard_data["faculty_stats"])


def display_metric_cards(metrics):
    """Display metric cards in dashboard"""
    st.markdown(
        f"""
        <div class="stats-container">
            <div class="stat-card">
                <h3>{metrics['total_requests']}</h3>
                <p>Total Requests</p>
            </div>
            <div class="stat-card">
                <h3>{metrics['pending_count']}</h3>
                <p>Pending Approval</p>
            </div>
            <div class="stat-card">
                <h3>{metrics['approved_count']}</h3>
                <p>Approved Requests</p>
            </div>
            <div class="stat-card">
                <h3>AED {metrics['total_cost']:,.2f}</h3>
                <p>Total Cost</p>
            </div>
        </div>
    """,
        unsafe_allow_html=True,
    )


def create_monthly_trend_config():
    """Special config for monthly trend chart with dark text"""
    return {
        "template": "plotly_white",
        "height": 400,
        "margin": dict(l=40, r=40, t=60, b=40),
        "font": dict(
            family="Inter",
            size=14,
            color="#18141D",  # Keep dark text for this chart
        ),
        "paper_bgcolor": "white",
        "plot_bgcolor": "white",
        "showlegend": True,
        "legend": dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
    }


def create_monthly_trend_chart(monthly_stats):
    """Create monthly trends visualization"""
    df = pd.DataFrame(monthly_stats)
    if not df.empty:
        df["month"] = pd.to_datetime(df["month"])
        fig = go.Figure()

        metrics = ["total_requests", "approved_requests", "pending_requests"]
        colors = ["#3C313D", "#556B2F", "#8B7355"]

        for metric, color in zip(metrics, colors):
            fig.add_trace(
                go.Scatter(
                    x=df["month"],
                    y=df[metric],
                    name=metric.replace("_", " ").title(),
                    line=dict(color=color, width=2),
                )
            )

        fig.update_layout(
            **create_monthly_trend_config(),  # Use special config for this chart
            title="Monthly Request Trends",
            xaxis_title="Month",
            yaxis_title="Number of Requests",
        )
        st.plotly_chart(fig, use_container_width=True)


def create_status_distribution_chart(status_stats):
    """Create status distribution pie chart"""
    df = pd.DataFrame(status_stats)
    if not df.empty:
        fig = px.pie(
            df,
            values="count",
            names="status",
            title="Request Status Distribution",
            color_discrete_sequence=["#556B2F", "#8B7355", "#8B4513"],
        )
        config = create_chart_config(is_monthly_trend=False)
        fig.update_layout(
            **config,
            title_font=dict(color="white", size=16),
        )
        fig.update_traces(
            textfont_color="white",
            hovertemplate="<b>%{label}</b><br>Count: %{value}<extra></extra>",
        )
        st.plotly_chart(fig, use_container_width=True)


def create_destination_chart(destination_stats):
    """Create destination analysis bar chart"""
    df = pd.DataFrame(destination_stats)
    if not df.empty:
        fig = px.bar(
            df,
            x="destination",
            y="visit_count",
            title="Top Destinations by Visit Count",
            color="total_cost",
            color_continuous_scale="Purples",
        )

        config = create_chart_config()
        fig.update_layout(
            **{
                k: v for k, v in config.items() if k not in ["xaxis", "yaxis"]
            },  # Remove axis settings from config
            title_font=dict(color="white", size=16),
            xaxis=dict(
                title="Destination",
                tickfont=dict(color="white"),
                title_font=dict(color="white"),
            ),
            yaxis=dict(
                title="Number of Visits",
                tickfont=dict(color="white"),
                title_font=dict(color="white"),
            ),
            coloraxis_colorbar=dict(
                title="Total Cost (AED)",
                tickfont=dict(color="white"),
                title_font=dict(color="white"),
            ),
        )

        fig.update_xaxes(tickangle=45)
        st.plotly_chart(fig, use_container_width=True)


def create_faculty_analysis_chart(faculty_stats):
    """Create faculty travel analysis chart"""
    df = pd.DataFrame(faculty_stats)
    if not df.empty:
        fig = go.Figure()

        fig.add_trace(
            go.Bar(
                x=df["name"],
                y=df["travel_count"],
                name="Travel Count",
                marker_color="#556B2F",
            )
        )

        fig.add_trace(
            go.Bar(
                x=df["name"],
                y=df["total_spent"],
                name="Total Spent (AED)",
                marker_color="#8B7355",
                yaxis="y2",
            )
        )
        custom_config = {
            "template": "plotly_white",
            "height": 400,
            "margin": dict(l=40, r=40, t=60, b=40),
            "font": dict(family="Inter", size=14, color="white"),
            "paper_bgcolor": "#3C313D",
            "plot_bgcolor": "#3C313D",
            "showlegend": True,
        }

        fig.update_layout(
            **custom_config,
            title=dict(
                text="Faculty Travel Analysis", font=dict(color="white", size=16)
            ),
            xaxis=dict(
                title="Faculty Member",
                tickfont=dict(color="white"),
                title_font=dict(color="white"),
            ),
            yaxis=dict(
                title="Travel Count",
                tickfont=dict(color="white"),
                title_font=dict(color="white"),
            ),
            yaxis2=dict(
                title="Total Spent (AED)",
                overlaying="y",
                side="right",
                tickfont=dict(color="white"),
                title_font=dict(color="white"),
            ),
            barmode="group",
            legend=dict(
                font=dict(color="white"),
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1,
            ),
        )
        fig.update_xaxes(tickangle=45)
        st.plotly_chart(fig, use_container_width=True)

def get_autofill(research_paper):
    try:
        st.session_state.research_paper = research_paper
        trimmed_text = extract_text_from_file(research_paper)[:7500]
        prompt = f"Please extract the following fields from \"input_text\" and output a single JSON object with keys \"title\",\"abstract_keywords\",\"purpose\",\"index_type\",\"country\",\"city\",\"from_date\",\"end_date\",\"per_diem\",\"registration_fee\",\"visa_fee\",\"conference_link\": \"title\" is the full conference title; \"abstract_keywords\" is an array of keywords from the abstract, if you couldn't find the abstract keywords just try to come up with some related to the conference; \"purpose\" is either \"Attending a Conference\" or \"Presenting at a Conference\"; \"index_type\" is one of [\"Scopus\",\"IEEE\",\"Web of Science\",\"PubMed\",\"MEDLINE\",\"ACM\"] if you did not find the index return Scopus; \"country\" is one of [\"USA\",\"Canada\",\"Germany\"] If you did not find the country return USA; \"city\" is one of [\"New York\",\"Los Angeles\",\"Chicago\",\"Toronto\",\"Vancouver\",\"Montreal\",\"Berlin\",\"Munich\",\"Hamburg\"] If you didn't find the city return the first city related to its country; \"from_date\" and \"end_date\" are in YYYY-MM-DD (ISO 8601); \"per_diem\",\"registration_fee\",\"visa_fee\" are numeric, defaulting to 100 if missing; \"conference_link\" is the URL; unspecified index_type, country, or city must be null; missing fees default to 100; output must be strictly valid JSON with exactly these keys. Here is the text to process: \"{trimmed_text}\""
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": prompt}
                ]
                )
        json_recommendation_data = json.loads(response.choices[0].message.content)
        # try:
        st.session_state.conference_title = json_recommendation_data.get("title")
        st.session_state.conference_abstract = json_recommendation_data.get("abstract_keywords")
        # st.write(json_recommendation_data.get("title"))
        if json_recommendation_data.get("purpose") in purpose:
            st.session_state.purpose = purpose.index(json_recommendation_data.get("purpose"))
        else:
            st.session_state.purpose = 0
        if json_recommendation_data.get("index_type") in indexes:
            st.session_state.index_type = indexes.index(json_recommendation_data.get("index_type"))
        else:
            st.session_state.index_type = 0
        select_country = list(location.keys())
        if json_recommendation_data.get("country") in location:
            st.session_state.country = json_recommendation_data.get("country")
            st.session_state.country_index = select_country.index(json_recommendation_data.get("country"))
        else:
            st.session_state.country = 0
        select_city = location[st.session_state.country]
        # st.write(select_city)
        if json_recommendation_data.get("city") in location[st.session_state.country]:
            st.session_state.city = select_city.index(json_recommendation_data.get("city"))
        else:
            st.session_state.city = 0
        if json_recommendation_data.get("from_date") is not None:
            st.session_state.from_date = datetime.strptime(json_recommendation_data["from_date"], "%Y-%m-%d").date()
        else:
            st.session_state.from_date = date.today()
        if json_recommendation_data.get("end_date") is not None:
            st.session_state.end_date = datetime.strptime(json_recommendation_data["end_date"], "%Y-%m-%d").date()
        else:
            st.session_state.end_date = date.today()
        if json_recommendation_data.get("per_diem"):
            st.session_state.per_diem = int(json_recommendation_data.get("per_diem"))
        else:
            st.session_state.per_diem = 100
        if json_recommendation_data.get("registration_fee"):
            st.session_state.registration_fee = int(json_recommendation_data.get("registration_fee"))
        else:
            st.session_state.registration_fee = 100
        if json_recommendation_data.get("visa_fee"):
            st.session_state.visa_fee = int(json_recommendation_data.get("visa_fee"))
        else:
            st.session_state.visa_fee = 100
        if json_recommendation_data.get("conference_link"):
            st.session_state.conference_link = json_recommendation_data.get("conference_link")
        else:
            st.session_state.conference_link = ''
    except Exception as e:
        st.error(e)
# Main application page
def main_app():
    # Center Navigation header (place this at the start of the function, right after the main header)
    st.sidebar.markdown(
        '<h2 style="text-align: center; color: var(--text-on-primary);">Navigation</h2>',
        unsafe_allow_html=True,
    )

    # Main Header
    st.markdown(
        """
        <div class="main-header">
            <h1>Conference Travel System</h1>
            <p>Manage your conference travel requests efficiently</p>
        </div>
    """,
        unsafe_allow_html=True,
    )

    # Welcome message with stats
    st.markdown(
        f"""
        <div class="custom-card">
            <h2>Welcome, {st.session_state.user_name}! 👋</h2>
            <p>Access your travel management dashboard and start planning your next conference.</p>
        </div>
    """,
        unsafe_allow_html=True,
    )

    # Stats cards (example for professors)
    if st.session_state.user_role == "professor":
        counts, _ = fetch_user_request_counts(st.session_state.logged_in_user)
        st.markdown(
            f"""
            <div class="stats-container">
                <div class="stat-card">
                    <h3>{counts['pending_count']}</h3>
                    <p>Pending Requests</p>
                </div>
                <div class="stat-card">
                    <h3>{counts['approved_count']}</h3>
                    <p>Approved Requests</p>
                </div>
                <div class="stat-card">
                    <h3>{counts['total_count']}</h3>
                    <p>Total Requests</p>
                </div>
            </div>
        """,
            unsafe_allow_html=True,
        )

    # Rest of your main_app code...

    user_role = st.session_state.user_role

    # Professor role with Home page and sidebar options
    if user_role == "professor":
        if "selected_tab" not in st.session_state:
            st.session_state.selected_tab = "home"

        # Sidebar navigation for professors
        if st.sidebar.button("Home"):
            st.session_state.selected_tab = "home"
        if st.sidebar.button("Conference Recommendation"):
            st.session_state.selected_tab = "recommendation"
        if st.sidebar.button("Paper Formatting and Evaluation"):
            st.session_state.selected_tab = "paper_formatting"
        if st.sidebar.button("Travel Form and Request Tracker"):
            st.session_state.selected_tab = "form_tracker"
        st.sidebar.markdown("<br>" * 2, unsafe_allow_html=True)

        # Add logout button at bottom of sidebar
        if st.sidebar.button("🚪 Logout", key="prof_logout_button"):
            logout()
            st.rerun()

        # Home Page content
        if st.session_state.selected_tab == "home":
            # Welcoming Header

            # Brief Introduction
            st.markdown("""
            ### Plan your conference travels effortlessly!
            Use this system to:
            - Submit your travel requests.
            - Track approvals.
            - Manage your budget and expenses.
            
            Navigate through the tabs on the sidebar to explore the system's features.
            """)

            # Decorative Separator
            st.markdown("---")

            # Call-to-Action Message
            st.markdown("""
            #### Ready to get started?
            - Head over to the **Travel Form & Request Tracker** tab to submit your travel details.
            - Check out the **Approval** tab for pending or approved requests.
            """)

            # Footer with a Tip or Closing Note
            st.markdown("""
            ---
            *Tip: Ensure your travel dates and budget align with system guidelines.*
            """)

        # Travel Form and Request Tracker
        elif st.session_state.selected_tab == "form_tracker":
            tabs = st.tabs(["Travel Form", "Request Tracker"])

            # Request Tracker Tab
            with tabs[1]:
                st.header("Request Summary")
                counts, user_requests = fetch_user_request_counts(
                    st.session_state.logged_in_user
                )
                st.write(f"**Pending Requests:** {counts['pending_count']}")
                st.write(f"**Approved Requests:** {counts['approved_count']}")
                st.write(f"**Total Requests:** {counts['total_count']}")

                user_requests_df = pd.DataFrame(user_requests)
                if not user_requests_df.empty:
                    st.dataframe(user_requests_df)
                else:
                    st.write("No requests found.")

                # Travel Form Tab
                with tabs[0]:
                    st.header("Submit Travel Details")
                    already_traveled = has_traveled_this_year(
                        st.session_state.logged_in_user
                    )
                    if already_traveled:
                        st.error(
                            "You have already traveled this year and cannot submit another travel request."
                        )
                    else:
                     # Initialize session state variables if they don't exist
                        if 'conference_title' not in st.session_state:
                            st.session_state.conference_title = ''
                        if 'conference_abstract' not in st.session_state:
                            st.session_state.conference_abstract = ''
                        if 'purpose' not in st.session_state:
                            st.session_state.purpose = 0
                        if 'index_type' not in st.session_state:
                            st.session_state.index_type = 0
                        if "country" not in st.session_state:
                            st.session_state.country = 0
                        if "country_index" not in st.session_state:
                            st.session_state.country_index = 0
                        if 'city' not in st.session_state:
                            st.session_state.city = 0
                        if 'from_date' not in st.session_state:
                            st.session_state.from_date = date.today()
                        if 'end_date' not in st.session_state:
                            st.session_state.end_date = date.today()
                        if 'per_diem' not in st.session_state:
                            st.session_state.per_diem = 100
                        if 'registration_fee' not in st.session_state:
                            st.session_state.registration_fee = 100
                        if 'visa_fee' not in st.session_state:
                            st.session_state.visa_fee = 100
                        if 'conference_link' not in st.session_state:
                            st.session_state.conference_link = ''
                        if 'research_paper' not in st.session_state:
                            st.session_state.research_paper = None
                        research_paper = st.file_uploader("Upload research paper to autofill the feilds")
                        if st.button("Submit Research Paper"):
                            get_autofill(research_paper)
                        # Conference Information
                        st.markdown("#### Conference Information")
                        conference_name = st.text_input(
                            "Name of Conference",
                            help="Enter the full name of the conference",
                            value=st.session_state.conference_title,
                        )

                        # Purpose and Index Type
                        col1, col2 = st.columns(2)
                        with col1:
                            purpose_of_attending = st.selectbox(
                                "Purpose of Attending",
                                purpose,
                                index=int(st.session_state.purpose)
                            )
                        with col2:
                            index_type = st.selectbox(
                                "Index Type",
                                indexes,
                                index=int(st.session_state.index_type)
                            )

                        # Location Details
                        st.markdown("#### Location Details")
                        col3, col4 = st.columns(2)
                        with col3:
                            country = st.selectbox(
                                "Choose Country", options=location.keys(),
                                index=int(st.session_state.country_index)
                            )
                        with col4:
                            city = st.selectbox(
                                "Choose City", options=location[country],
                                index=st.session_state.city
                            )

                        # Travel Dates
                        st.markdown("#### Travel Dates")
                        col5, col6 = st.columns(2)
                        with col5:
                            date_from = st.date_input("From Date", value=st.session_state.from_date)
                        with col6:
                            date_to = st.date_input("To Date", value=st.session_state.end_date)

                        # Date validation: Check for invalid date range and restricted periods
                        allowed, message = is_within_restricted_dates(
                            date_from, date_to
                        )
                        if not allowed:
                            st.warning(
                                message
                            )  # Warning message if dates overlap with restricted periods

                        if date_from > date_to:
                            st.error("The 'From Date' cannot be after the 'To Date'.")

                        # Financial Details
                        st.markdown("#### Financial Details (in AED)")
                        col7, col8, col9 = st.columns(3)
                        with col7:
                            per_diem = st.number_input(
                                "Per Diem (AED)", min_value=100, step=100, format="%d", value=st.session_state.per_diem
                            )
                        with col8:
                            registration_fee = st.number_input(
                                "Registration Fee (AED)",
                                min_value=100,
                                step=100,
                                format="%d",
                                value=st.session_state.registration_fee
                            )
                        with col9:
                            visa_fee = st.number_input(
                                "Visa Fee (AED)", min_value=100, step=100, format="%d", value=st.session_state.visa_fee
                            )

                        # Supporting Documents
                        st.markdown("#### Supporting Documents")
                        acceptance_letter = st.file_uploader(
                            "Upload Acceptance Letter PDF", type=["pdf"]
                        )
                        if acceptance_letter:
                            is_valid, message = validate_file_upload(acceptance_letter)
                            if not is_valid:
                                st.error(message)
                        if st.session_state.research_paper is not None:
                            st.info("Research Paper has been already uploaded!")
                            research_paper = st.session_state.research_paper
                        else:
                            research_paper = st.file_uploader(
                                   "Upload Research Paper PDF", type=["pdf"]
                               )
                        conference_url = st.text_input(
                            "Enter Conference URL", value=st.session_state.conference_link
                        )

                        # Initialize variables
                        url_feedback = None
                        url_summary = None

                        if conference_url:
                            url_feedback = process_url(conference_url)

                        if url_feedback:  # Ensure url_feedback is not None
                            if "Valid URL" in url_feedback:
                                url_summary = url_feedback.replace(
                                    "Valid URL. Summary:", ""
                                ).strip()
                            elif "Invalid URL" in url_feedback:
                                st.warning(url_feedback)
                        else:
                            url_summary = None  # Handle the case where process_url returns None
                                                # Submit Button
                        submit_button = st.button("Submit Travel Details")

                        if submit_button and not st.session_state.get("submitted", False):
                            # Field validation
                            if not conference_name:
                                st.error("Please enter the name of the conference.")
                            elif index_type == "None":
                                st.error("You must select a valid index type. 'None' is not allowed.")
                                st.session_state["submitted"] = False
                            elif not conference_url:
                                st.error("Conference URL is a required field.")
                            elif not validate_url(conference_url):
                                st.error("Invalid URL. Please check the conference URL and try again.")
                            elif not acceptance_letter or not research_paper:
                                st.error("Please upload both the Acceptance Letter and Research Paper.")
                            else:
                                # Validate travel dates
                                allowed, message = is_within_restricted_dates(date_from, date_to)
                                if not allowed:
                                    st.error(message)
                                    st.session_state["submitted"] = False
                                elif date_from > date_to:
                                    st.error("The 'From Date' cannot be after the 'To Date'.")
                                    st.session_state["submitted"] = False
                                else:
                                    try:
                                        # Extract text from uploaded documents
                                        research_paper_text = extract_text_from_file(research_paper)
                                        acceptance_letter_text = extract_text_from_file(acceptance_letter)

                                        if research_paper_text and acceptance_letter_text:
                                            with st.spinner("Validating documents with AI..."):
                                                validation = validate_with_gpt(
                                                    conference_name,
                                                    research_paper_text[:700],
                                                    acceptance_letter_text[:700],
                                                    index_type,
                                                    conference_url
                                                )

                                            # Handle index verification
                                            if validation.get("valid"):
                                                if validation.get("index_verified") is False:
                                                    st.error(f"Conference index mismatch for {index_type}")
                                                    st.session_state["submitted"] = False
                                                elif validation.get("index_verified") is None:
                                                    st.warning("Index verification inconclusive - requires manual review")
                                                    st.session_state["submitted"] = False
                                                else:
                                                    # Budget check
                                                    total_cost = per_diem + registration_fee + visa_fee
                                                    available_budget = fetch_available_budget()

                                                    if total_cost > available_budget:
                                                        st.error(f"Insufficient budget! Total: AED {total_cost}, Available: AED {available_budget}")
                                                        st.session_state["submitted"] = False
                                                    else:
                                                        # Process the URL
                                                        url_summary = process_url(conference_url)

                                                        # Insert request into the database
                                                        data = (
                                                            st.session_state.logged_in_user, conference_name, purpose_of_attending,
                                                            index_type, country, city, date_from, date_to, per_diem,
                                                            registration_fee, visa_fee, conference_url, url_summary
                                                        )

                                                        request_id = insert_travel_request(
                                                            data,
                                                            acceptance_letter,
                                                            conference_url,
                                                            research_paper
                                                        )

                                                        if request_id == "duplicate":
                                                            st.error("Duplicate request found. Submission disallowed.")
                                                        else:
                                                            insert_uploaded_files(request_id, acceptance_letter, research_paper)
                                                            st.success("Request submitted successfully!")
                                                            st.session_state["submitted"] = True
                                            else:
                                                st.error(f"Request Rejected: {validation['message']}")
                                                for issue in validation.get("issues", []):
                                                    st.write(f"- {issue}")
                                                st.session_state["submitted"] = False
                                        else:
                                            st.error("Document processing failed.")
                                            st.session_state["submitted"] = False

                                    except Exception as e:
                                        st.error(f"Submission error: {str(e)}")
                                        st.session_state["submitted"] = False
                                        logging.error(f"Submission failed: {str(e)}")
        # Conference Recommendation page placeholder
        elif st.session_state.selected_tab == "recommendation":
            st.markdown(
                """
        <style>
        /* Input field styling for recommendation form */
        div[data-baseweb="input"] input,
        div[data-baseweb="textarea"] textarea {
            background-color: transparent !important;
            /*color: white !important;*/
            border: 2px solid var(--border) !important;
            padding: 0.75rem 1rem !important;
            font-size: 1rem !important;
            border-radius: 8px !important;
            width: 100% !important;
            /*caret-color: white !important;*/
        }


        /* Placeholder text */
        div[data-baseweb="input"] input::placeholder,
        div[data-baseweb="textarea"] textarea::placeholder {
            color: var(--text-secondary) !important;
            opacity: 0.7;
        }

        /* Focus state */
        div[data-baseweb="input"] input:focus,
        div[data-baseweb="textarea"] textarea:focus {
            border-color: var(--primary) !important;
            box-shadow: 0 0 0 2px rgba(60, 49, 61, 0.1) !important;
        }
        </style>
    """,
                unsafe_allow_html=True,
            )
            # Input fields
            recommendations = None
            research_paper_recommendation = st.file_uploader(
                "Upload Research Paper PDF (Optional)"
            )
            # Initialize session state variables if they don't exist
            if 'conference_title' not in st.session_state:
                st.session_state.conference_title = ''
            if 'conference_abstract' not in st.session_state:
                st.session_state.conference_abstract = ''
            if st.button("Submit Research Paper"):
                get_autofill(research_paper_recommendation)
            conference_title = st.text_input(
            "Abstract Title",
            placeholder="Enter the title of your research or conference topic",
            value=st.session_state.conference_title
            )
            conference_abstract = st.text_area(
            "Abstract", placeholder="Enter the abstract for your research or ideas",
            value=st.session_state.conference_abstract
            )
            # Submit button
            if st.button("Submit"):
                if conference_title and conference_abstract:
                    user_input = (
                        f"Title: {conference_title}\nAbstract: {conference_abstract}",
                    )
                    try:
                        gClient = genai.Client(api_key=os.getenv("GOOGLE_AI_API_KEY"))
                        response = gClient.models.generate_content(
                            model='gemini-2.0-flash-001',
                            contents=f"You are an expert in academic conferences. Given a research title and abstract, search the web and recommend up to 5 relevant academic conferences that will take place in 2025 onwards. For each recommendation, include the following details in JSON format: conference title, description (keep it short up to 50 words), location, index, exact event date in DD/MM/YYYY format, exact submission deadline date  in DD/MM/YYYY format and the conference link of the company that is hosting that event. Return the JSON as a list of objects with keys: 'conference_title', 'description', 'location', 'index', 'event_date', 'submission_deadline', 'conference link'. Here it the title: {conference_title} and the abstract: {conference_abstract}. Make sure you put accurate dates and don't write something like: Typically in some Month, make sure it is in DD/MM/YYYY format and make sure it is accurate from the confernce website.")
                        config=types.GenerateContentConfig(
                            tools=[types.Tool(
                                google_search=types.GoogleSearchRetrieval
                                )]
                                )
                        match = re.search(r"```json(.*?)```", response.text, re.DOTALL)
                        if match:
                            json_data = match.group(1).strip()
                            res = json.loads(json_data)
                            dataframe = pd.DataFrame(res)
                            st.subheader("Confrences")
                            st.dataframe(dataframe, column_config={"conference link": st.column_config.LinkColumn()})
                        else:
                            st.error("No confrences found")
                    except Exception as e:
                      st.error(f"Error fetching recommendations: {e}")
                else:
                    st.warning("Please fill in both the title and abstract.")
        elif st.session_state.selected_tab == "paper_formatting":
            st.markdown(
                """
        <style>
        /* Input field styling for paper formatting form */
        div[data-baseweb="input"] input,
        div[data-baseweb="textarea"] textarea {
            background-color: transparent !important;
            /*color: var(--text-on-primary) !important;*/
            border: 2px solid var(--border) !important;
            padding: 0.75rem 1rem !important;
            font-size: 1rem !important;
            border-radius: 8px !important;
            width: 100% !important;
            caret-color: var(--primary) !important;
        }

        /* Placeholder text */
        div[data-baseweb="input"] input::placeholder,
        div[data-baseweb="textarea"] textarea::placeholder {
            color: var(--text-secondary) !important;
            opacity: 0.7;
        }

        /* Focus state */
        div[data-baseweb="input"] input:focus,
        div[data-baseweb="textarea"] textarea:focus {
            border-color: var(--primary) !important;
            box-shadow: 0 0 0 2px rgba(60, 49, 61, 0.1) !important;
        }
        </style>
    """,
                unsafe_allow_html=True,
            )
            # Input fields
            paper_type = st.selectbox(
                "Type of Paper",
                [
                    "IEEE Paper",
                    "ACM Paper",
                    "Springer Paper",
                    "Elsevier Paper",
                ],
            )
            
            # Option to choose between file upload or text area
            option = st.radio("Choose input method", ("Upload a DOCX or PDF file", "Enter text manually"))

            if option == "Upload a DOCX or PDF file":
                uploaded_file = st.file_uploader("Upload a DOCX or PDF file", type=["docx", "pdf"])
                paper_text = ""
            elif option == "Enter text manually":
                paper_text = st.text_area(
                    "Enter your text here",
                    placeholder="Enter the text you want to convert to LaTeX",
                )
                uploaded_file = None

            # Extract text from uploaded file if any
            if uploaded_file is not None:
                if uploaded_file.type == "application/pdf":
                    # Extract text from PDF
                    from PyPDF2 import PdfReader
                    reader = PdfReader(uploaded_file)
                    paper_text = "\n".join(page.extract_text() for page in reader.pages)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # Extract text from DOCX
                    import docx
                    doc = docx.Document(uploaded_file)
                    paper_text = "\n".join(para.text for para in doc.paragraphs)

            # Submit button
            if st.button("Convert to LaTeX"):
                if paper_type and paper_text:
                    user_input = (
                        f"Type of Paper: {paper_type}\nText: {paper_text}"
                    )

                    # Display a loading indicator
                    with st.spinner('Converting to LaTeX...'):
                        # Query OpenAI for strengths and weaknesses
                        try:
                            feedback_response = client.chat.completions.create(
                                model="gpt-4",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": (
                                            "You are an expert in academic paper evaluation. Provide strengths and weaknesses of the paper based on the text."
                                        ),
                                    },
                                    {"role": "user", "content": paper_text},
                                ],
                            )
                            feedback_output = feedback_response.choices[0].message.content.strip()
                            st.subheader("Strengths and Weaknesses:")
                            st.write(feedback_output)

                            # Query OpenAI for LaTeX conversion
                            response = client.chat.completions.create(
                                model="gpt-3.5-turbo",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": (
                                            "You are an expert in LaTeX formatting. Given the type of paper and the text, "
                                            "convert the text into LaTeX format based on the type of paper. "
                                            "Only output LaTeX."
                                        ),
                                    },
                                    {"role": "user", "content": user_input},
                                ],
                            )
                            # Parse the LaTeX response
                            latex_output = response.choices[
                                0
                            ].message.content.strip()                            

                            # Create a temporary .tex file
                            with open("temp.tex", "w") as f:
                                f.write(latex_output)
                            
                            # Write LaTeX to temporary file
                            with open("temp.tex", "w") as f:
                                f.write(latex_output)

                            st.subheader("LaTeX Output:")
                            st.code(latex_output, language="latex")

                        except Exception as e:
                            st.error(f"Error converting to LaTeX: {e}")
                else:
                    st.warning("Please select the type of paper and enter the text.")
    elif user_role == "accountant":
        st.header("Budget Manager")
        st.markdown("#### Set or Adjust Total Available Budget (AED)")
        st.sidebar.markdown("<br>" * 2, unsafe_allow_html=True)

        # Add logout button at bottom of sidebar
        if st.sidebar.button("🚪 Logout", key="prof_logout_button"):
            logout()
            st.rerun()

        # Define callback functions for budget management
        def confirm_budget_callback():
            set_budget(st.session_state["new_budget"])  # Set the budget
            st.success("Budget set or adjusted successfully!")
            st.session_state.refresh_budget = (
                not st.session_state.refresh_budget
            )  # Trigger refresh
            st.session_state.pop("confirm_budget", None)
            st.session_state.pop("new_budget", None)

        def cancel_budget_callback():
            st.info("Budget setting canceled.")
            st.session_state.pop("confirm_budget", None)
            st.session_state.pop("new_budget", None)

        # Main form to set the budget with confirmation
        with st.form("Budget Form"):
            available_budget = fetch_available_budget()
            new_budget = st.number_input(
                "Total Available Budget (AED)",
                min_value=0.0,
                value=available_budget,
                format="%.2f",
            )
            submit_budget_button = st.form_submit_button("Set Budget")
            if submit_budget_button:
                st.session_state["new_budget"] = new_budget
                st.session_state["confirm_budget"] = True

        # If the form has been submitted and confirmation is required
        if st.session_state.get("confirm_budget", False):
            st.warning(
                f"Are you sure you want to set the budget to AED {st.session_state['new_budget']:,.2f}?"
            )
            st.button("Confirm Budget Setting", on_click=confirm_budget_callback)
            st.button("Cancel", on_click=cancel_budget_callback)

        # Display current available budget
        st.markdown("#### Current Available Budget")
        st.write(f"**AED {fetch_available_budget():,.2f}**")

        # Display the budget adjustment and deduction history
        st.markdown("#### Budget Adjustment and Deduction History")
        budget_history_data = fetch_budget_history()
        history_df = pd.DataFrame(budget_history_data)
        if not history_df.empty:
            st.table(history_df)
        else:
            st.write("No budget adjustment history available.")

    # Approval Tab for Approval Users
    elif user_role == "approval":
        st.sidebar.header("Approval Section")

        tab_selected = st.sidebar.radio(
            "Select Section",
            ["📝 Pending Requests", "📊 Dashboard"],
            key="approval_tabs",
        )

        st.sidebar.markdown("<br>" * 2, unsafe_allow_html=True)
        if st.sidebar.button("🚪 Logout", key="logout_button"):
            logout()
            st.rerun()

        if tab_selected == "📝 Pending Requests":
            display_pending_requests()

        elif tab_selected == "📊 Dashboard":
            display_dashboard()

        # Key Metrics Cards
        connection = get_db_connection()
        cursor = connection.cursor(dictionary=True)

        # Fetch key metrics
        cursor.execute("""
            SELECT 
                COUNT(*) as total_requests,
                SUM(CASE WHEN status = 'pending' THEN 1 ELSE 0 END) as pending_count,
                SUM(CASE WHEN status = 'approved' THEN 1 ELSE 0 END) as approved_count,
                SUM(per_diem + registration_fee + visa_fee) as total_cost
            FROM requests
        """)
        # Display metrics in modern card


# Callback function to delete the request
def delete_request_callback():
    request_id = st.session_state["delete_request_id"]
    delete_request(request_id)
    st.success(f"Request ID {request_id} deleted successfully.")
    del st.session_state["delete_request_id"]  # Clear after deletion


def delete_request(request_id):
    connection = get_db_connection()
    cursor = connection.cursor()

    # Delete associated files first to satisfy the foreign key constraint
    cursor.execute("DELETE FROM UploadedFiles WHERE request_id = %s", (request_id,))
    cursor.execute("DELETE FROM Requests WHERE request_id = %s", (request_id,))

    connection.commit()
    cursor.close()
    connection.close()
    st.success(f"Request ID {request_id} deleted successfully.")


# Logout function
def logout():
    """Enhanced secure logout with activity logging"""
    if st.session_state.logged_in_user:
        # Log the logout action
        user_id = st.session_state.logged_in_user
        logging.info(f"User {user_id} logged out")

        # Record logout in activity log
        record_user_activity(
            user_id, "logout", {"session_id": st.session_state.session_id}
        )

    # Clear all session state
    for key in list(st.session_state.keys()):
        if key != "session_id":  # Keep session ID for tracking
            del st.session_state[key]

    # Reset basic session variables
    st.session_state.page = "login"
    st.session_state.logged_in_user = None
    st.session_state.user_role = None
    st.session_state.last_activity = datetime.now()
    st.session_state.login_attempts = 0
    st.session_state.logout_requested = True


# Add this modern CSS styling
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    
    /* Main Styles */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    
    /* Cards */
    div.stButton > button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        background-color: #4CAF50;
        /*color: white;*/
        font-weight: 600;
        border: none;
        transition: all 0.3s ease;
    }
    div.stButton > button:hover {
        background-color: #45a049;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        transform: translateY(-2px);
    }
    
    /* Sidebar */
    section[data-testid="stSidebar"] {
       /* background-color: #1a1a1a;*/
        padding: 2rem 0;
    }
    section[data-testid="stSidebar"] .stButton > button {
        background-color: transparent;
        /*color: #ffffff;*/
        text-align: left;
        padding: 1rem 2rem;
        border-radius: 0;
        border-left: 4px solid transparent;
    }
    section[data-testid="stSidebar"] .stButton > button:hover {
        /*background-color: rgba(255,255,255,0.1);*/
        border-left: 4px solid #4CAF50;
        transform: none;
    }
    
    /* Headers */
    h1, h2, h3 {
        color: #1a1a1a;
        font-weight: 700;
    }
    
    /* Cards/Containers */
    div.stDataFrame {
        border: 1px solid #e0e0e0;
        border-radius: 10px;
        padding: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    /* Form inputs */
    div.stTextInput > div > div > input {
        border-radius: 8px;
        border: 2px solid #e0e0e0;
        padding: 0.5rem 1rem;
        font-size: 16px;
    }
    div.stTextInput > div > div > input:focus {
        border-color: #4CAF50;
        box-shadow: 0 0 0 2px rgba(76,175,80,0.2);
    }
    
    /* Select boxes */
    div.stSelectbox > div > div > select {
        border-radius: 8px;
        border: 2px solid #e0e0e0;
    }
    
    /* Metrics */
    div[data-testid="stMetricValue"] {
        background-color: #f8f9fa;
        padding: 1rem 2rem;
        color: #666666;
        font-weight: 600;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        color: #4CAF50;
        border-bottom: 4px solid #4CAF50;
    }
    
    /* Custom Cards */
    .custom-card {
       /* background-color: white;*/
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        margin-bottom: 1rem;
    }
    
    /* Stats Container */
    .stats-container {
        display: flex;
        justify-content: space-between;
        gap: 1rem;
        margin: 1rem 0;
    }
    .stat-card {
       /* background-color: white;*/
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        flex: 1;
        text-align: center;
    }
    
    /* Header with gradient */
    .main-header {
        background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
        /*color: white;*/
        padding: 2rem;
        border-radius: 10px;
        margin-
    /* Radio button text color fix */
    .stRadio [data-testid="stMarkdownContainer"] p {
        /*color: white !important;*/
    }
    
    /* Make radio button itself more visible */
    .stRadio input[type="radio"] {
        /*border-color: white !important;*/
    }
    
    /* Selected radio button text emphasis */
    .stRadio [aria-checked="true"] [data-testid="stMarkdownContainer"] p {
        font-weight: 600;
        /*color: #4CAF50 !important;*/
    }
    
    /* Sidebar radio group styling */
    [data-testid="stSidebar"] .stRadio {
        background: transparent;
        padding: 0.5rem 1rem;
        margin: 0.5rem 0;
    }
    
    /* ... rest of your existing styles ... */
    </style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
    <style>
    /* Modern Color Palette */
    :root {
        --primary: #2E7D32;
        --primary-light: #4CAF50;
        --primary-dark: #1B5E20;
        --secondary: #1A237E;
        --accent: #FFB300;
        --success: #43A047;
        --warning: #FFA726;
        --error: #E53935;
        --background: #FAFAFA;
        --surface: #FFFFFF;
        --text-primary: #212121;
        --text-secondary: #757575;
        --text-on-primary: #FFFFFF;
        --border: #E0E0E0;
    }

    /* Typography */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
        color: var(--text-primary);
        background-color: var(--background);
    }

    /* Main Header */
    .main-header {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
        /*color: white;*/
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }

    .main-header h1 {
        color: var(--text-on-primary);
        font-weight: 700;
        margin-bottom: 0.5rem;
        font-size: 2.2rem;
    }

    .main-header p {
        opacity: 0.9;
        font-size: 1.1rem;
    }

    /* Cards */
    .custom-card {
        background: var(--surface);
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        margin-bottom: 1.5rem;
        border: 1px solid rgba(0,0,0,0.05);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }

    .custom-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    }

    /* Stats Container */
    .stats-container {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1.5rem;
        margin: 1.5rem 0;
    }

    .stat-card {
        background: var(--surface);
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        text-align: center;
        border: 1px solid rgba(0,0,0,0.05);
        transition: all 0.3s ease;
    }

    .stat-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    }

    .stat-card h3 {
        color: var(--primary);
        font-size: 1.8rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }

    .stat-card p {
        color: var(--text-secondary);
        font-size: 1rem;
        font-weight: 500;
    }

    /* Buttons */
    div.stButton > button {
        width: 100%;
        padding: 0.75rem 1.5rem;
        border-radius: 8px;
        background: var(--primary);
        color: var(--text-on-primary);
        font-weight: 600;
        border: none;
        transition: all 0.3s ease;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        font-size: 0.9rem;
    }

    div.stButton > button:hover {
        background: var(--primary-dark);
        box-shadow: 0 4px 12px rgba(46,125,50,0.2);
        transform: translateY(-2px);
    }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: var(--primary-dark);
        padding: 2rem 0;
    }

    section[data-testid="stSidebar"] .stButton > button {
        background-color: transparent;
        color: var(--text-on-primary);
        text-align: left;
        padding: 1rem 2rem;
        border-radius: 0;
        border-left: 4px solid transparent;
        text-transform: none;
        letter-spacing: normal;
    }

    section[data-testid="stSidebar"] .stButton > button:hover {
        background-color: rgba(255,255,255,0.1);
        border-left: 4px solid var(--primary-light);
        transform: none;
    }

    /* Form Inputs */
    div.stTextInput > div > div > input,
    div.stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 2px solid var(--border);
        padding: 0.75rem 1rem;
        font-size: 1rem;
        transition: all 0.3s ease;
        background: var(--surface);
    }

    div.stTextInput > div > div > input:focus,
    div.stTextArea > div > div > textarea:focus {
        border-color: var(--primary);
        box-shadow: 0 0 0 2px rgba(46,125,50,0.1);
    }

    /* Select boxes */
    div.stSelectbox > div > div > select {
        border-radius: 8px;
        border: 2px solid var (--border);
        padding: 0.75rem 1rem;
    }

    /* Radio Buttons */
    .stRadio [data-testid="stMarkdownContainer"] p {
        /*color: var(--text-on-primary) !important;*/
    }

    .stRadio input[type="radio"] {
        /*border-color: var(--text-on-primary) !important;*/
    }

    .stRadio [aria-checked="true"] [data-testid="stMarkdownContainer"] p {
        /*color: var (--primary-light) !important;*/
        font-weight: 600;
    }

    /* Tables */
    .dataframe {
        border: none !important;
        border-collapse: separate;
        border-spacing: 0;
        width: 100%;
        margin: 1rem 0;
        background: var (--surface);
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
    }

    .dataframe th {
        background: var(--primary);
        color: var(--text-on-primary);
        padding: 1rem !important;
        font-weight: 600;
    }

    .dataframe td {
        padding: 1rem !important;
        border-top: 1px solid var(--border);
        color: var (--text-primary);
    }

    /* Charts */
    .js-plotly-plot {
        border-radius: 8px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        padding: 1rem;
        background: var (--surface);
    }

    /* Notifications */
    div[data-baseweb="notification"] {
        border-radius: 8px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }

    /* Expander */
    .streamlit-expander {
        border: none;
        border-radius: 8px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        background: var (--surface);
    }

    .streamlit-expander > div:first-child {
        border-radius: 8px 8px 0 0;
        padding: 1rem 1.5rem;
        background: var (--primary);
        color: var (--text-on-primary);
    }

    /* Recent Activity Items */
    .activity-item {
        padding: 1.2rem;
        border-radius: 8px;
        background: var (--surface);
        margin-bottom: 1rem;
        border: 1px solid var(--border);
        transition: all 0.3s ease;
    }

    .activity-item:hover {
        transform: translateX(4px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }

    /* Status Tags */
    .status-tag {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 500;
    }

    .status-approved {
        background: rgba(67,160,71,0.1);
        color: var(--success);
    }

    .status-pending {
        background: rgba(255,179,0,0.1);
        color: var(--warning);
    }

    .status-rejected {
        background: rgba(229,57,53,0.1);
        color: var(--error);
    }
    </style>
""",
    unsafe_allow_html=True,
)

# Update the color variables and text contrast
st.markdown(
    """
    <style>
    /* Modern Color Palette */
    :root {
        --primary: #18141D;
        --primary-light: #55535B;
        --primary-dark: #18141D;
        --secondary: #9D9BA2;
        --background: #EBEAEC;
        --surface: #FFFFFF;
        --text-primary: #18141D;
        --text-secondary: #55535B;
        --text-on-primary: #FFFFFF;
        --text-on-surface: #18141D;
        --border: #9D9BA2;          /* Border color */
        --success: #556B2F;         /* Subtle olive green */
        --warning: #8B7355;         /* Muted brown */
        --error: #8B4513;           /* Saddle brown */
    }

    /* Typography */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
        color: var(--text-primary);
        background-color: var(--background);
    }

    /* Main Header */
    .main-header {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
        /*color: white;*/
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }

    .main-header h1 {
        color: var(--text-on-primary);
        font-weight: 700;
        margin-bottom: 0.5rem;
        font-size: 2.2rem;
    }

    .main-header p {
        opacity: 0.9;
        font-size: 1.1rem;
    }

    /* Cards */
    .custom-card {
        background: var(--surface);
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        margin-bottom: 1.5rem;
        border: 1px solid rgba(0,0,0,0.05);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }

    .custom-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    }

    /* Stats Container */
    .stats-container {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1.5rem;
        margin: 1.5rem 0;
    }

    .stat-card {
        background: var(--surface);
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        text-align: center;
        border: 1px solid rgba(0,0,0,0.05);
        transition: all 0.3s ease;
    }

    .stat-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    }

    .stat-card h3 {
        color: var(--primary);
        font-size: 1.8rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }

    .stat-card p {
        color: var(--text-secondary);
        font-size: 1rem;
        font-weight: 500;
    }

    /* Buttons */
    div.stButton > button {
        width: 100%;
        padding: 0.75rem 1.5rem;
        border-radius: 8px;
        background: var(--primary);
        color: var(--text-on-primary);
        font-weight: 600;
        border: none;
        transition: all 0.3s ease;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        font-size: 0.9rem;
    }

    div.stButton > button:hover {
        background: var(--primary-dark);
        box-shadow: 0 4px 12px rgba(46,125,50,0.2);
        transform: translateY(-2px);
    }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: var(--primary-dark);
        padding: 2rem 0;
    }

    section[data-testid="stSidebar"] .stButton > button {
        background-color: transparent;
        color: var(--text-on-primary);
        text-align: left;
        padding: 1rem 2rem;
        border-radius: 0;
        border-left: 4px solid transparent;
        text-transform: none;
        letter-spacing: normal;
    }

    section[data-testid="stSidebar"] .stButton > button:hover {
        background-color: rgba(255,255,255,0.1);
        border-left: 4px solid var(--primary-light);
        transform: none;
    }

    /* Form Inputs */
    div.stTextInput > div > div > input,
    div.stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 2px solid var(--border);
        padding: 0.75rem 1rem;
        font-size: 1rem;
        transition: all 0.3s ease;
        background: var(--surface);
    }

    div.stTextInput > div > div > input:focus,
    div.stTextArea > div > div > textarea:focus {
        border-color: var(--primary);
        box-shadow: 0 0 0 2px rgba(46,125,50,0.1);
    }

    /* Select boxes */
    div.stSelectbox > div > div > select {
        border-radius: 8px;
        border: 2px solid var (--border);
        padding: 0.75rem 1rem;
    }

    /* Radio Buttons */
    .stRadio [data-testid="stMarkdownContainer"] p {
       /* color: var(--text-on-primary) !important;*/
    }

    .stRadio input[type="radio"] {
        /*border-color: var(--text-on-primary) !important;*/
    }

    .stRadio [aria-checked="true"] [data-testid="stMarkdownContainer"] p {
        /*color: var (--primary-light) !important;*/
        font-weight: 600;
    }

    /* Tables */
    .dataframe {
        border: none !important;
        border-collapse: separate;
        border-spacing: 0;
        width: 100%;
        margin: 1rem 0;
        background: var (--surface);
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
    }

    .dataframe th {
        background: var(--primary);
        color: var(--text-on-primary);
        padding: 1rem !important;
        font-weight: 600;
    }

    .dataframe td {
        padding: 1rem !important;
        border-top: 1px solid var(--border);
        color: var (--text-primary);
    }

    /* Charts */
    .js-plotly-plot {
        border-radius: 8px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        padding: 1rem;
        background: var (--surface);
    }

    /* Notifications */
    div[data-baseweb="notification"] {
        border-radius: 8px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }

    /* Expander */
    .streamlit-expander {
        border: none;
        border-radius: 8px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
        background: var (--surface);
    }

    .streamlit-expander > div:first-child {
        border-radius: 8px 8px 0 0;
        padding: 1rem 1.5rem;
        background: var (--primary);
        color: var (--text-on-primary);
    }

    /* Recent Activity Items */
    .activity-item {
        padding: 1.2rem;
        border-radius: 8px;
        background: var (--surface);
        margin-bottom: 1rem;
        border: 1px solid var(--border);
        transition: all 0.3s ease;
    }

    .activity-item:hover {
        transform: translateX(4px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }

    /* Status Tags */
    .status-tag {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 500;
    }

    .status-approved {
        background: rgba(67,160,71,0.1);
        color: var(--success);
    }

    .status-pending {
        background: rgba(255,179,0,0.1);
        color: var(--warning);
    }

    .status-rejected {
        background: rgba(229,57,53,0.1);
        color: var(--error);
    }
    </style>
""",
    unsafe_allow_html=True,
)

# Update the color palette with subtle, professional colors
st.markdown(
    """
    <style>
    :root {
        --primary: #3C313D;          /* Deep plum */
        --primary-light: #55535B;    /* Lighter plum */
        --primary-dark: #18141D;     /* Dark plum */
        --secondary: #9D9BA2;        /* Subtle grey */
        --background: #EBEAEC;       /* Light grey background */
        --surface: #FFFFFF;          /* White */
        --text-primary: #18141D;     /* Dark text */
        --text-secondary: #55535B;   /* Secondary text */
        --text-on-primary: #FFFFFF;  /* White text on dark */
        --text-on-surface: #18141D;  /* Dark text on light */
        --border: #9D9BA2;          /* Border color */
        --success: #556B2F;         /* Subtle olive green */
        --warning: #8B7355;         /* Muted brown */
        --error: #8B4513;           /* Saddle brown */
    }
    </style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
    <style>
    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2rem;
        border-bottom: 2px solid var(--border);
    }

    .stTabs [data-baseweb="tab"] {
        padding: 1rem 2rem;
        color: var(--text-secondary);
        font-weight: 500;
    }

    /* Selected tab styling */
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        color: var(--primary);
        border-bottom: 4px solid var(--primary);
        font-weight: 600;
    }

    /* Tab hover effect */
    .stTabs [data-baseweb="tab"]:hover {
        color: var(--primary-light);
        background: transparent;
    }

    /* Tab panel spacing */
    .stTabs [data-baseweb="tab-panel"] {
        padding-top: 2rem;
    }
    </style>
""",
    unsafe_allow_html=True,
)

# Update sidebar button text color
st.markdown(
    """
    <style>
    /* Sidebar button text color */
    section[data-testid="stSidebar"] .stButton > button {
        color: var(--text-on-primary) !important;
        background-color: transparent;
        text-align: left;
        padding: 1rem 2rem;
        border-radius: 0;
        border-left: 4px solid transparent;
        text-transform: none;
        letter-spacing: normal;
        width: 100%;
    }

    /* Selected/active state */
    section[data-testid="stSidebar"] .stButton > button:active,
    section[data-testid="stSidebar"] .stButton > button:focus {
        color: var(--text-on-primary) !important;
        background-color: rgba(255,255,255,0.1);
        border-left: 4px solid var(--primary-light);
    }

    /* Hover state */
    section[data-testid="stSidebar"] .stButton > button:hover {
        color: var(--text-on-primary) !important;
        background-color: rgba(255,255,255,0.1);
        border-left: 4px solid var(--primary-light);
    }
    </style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
    <style>
    /* Dark background, light text */
    section[data-testid="stSidebar"], 
    .main-header, 
    .streamlit-expander > div:first-child {
        color: var(--text-on-primary) !important;
    }
    
    /* Ensure sidebar buttons have white text */
    section[data-testid="stSidebar"] .stButton > button,
    section[data-testid="stSidebar"] .stButton > button:hover,
    section[data-testid="stSidebar"] .stButton > button:active,
    section[data-testid="stSidebar"] .stButton > button:focus {
        color: var(--text-on-primary) !important;
    }

    /* Light background, dark text */
    .custom-card,
    .stat-card,
    .activity-item,
    .dataframe td,
    .js-plotly-plot text {
        color: var(--text-primary) !important;
    }

    /* Fix chart text visibility */
    .js-plotly-plot .plotly .main-svg text {
        fill: var(--text-primary) !important;
    }

    /* Make Travel Form and Request Tracker text white */
    .stTabs [data-baseweb="tab"] {
        color: var(--text-on-primary) !important;
    }

    /* Selected tab styling */
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        color: var(--text-on-primary) !important;
        border-bottom-color: var(--text-on-primary);
    }
    </style>
""",
    unsafe_allow_html=True,
)


def reject_request(request_id):
    """Reject a travel request"""
    try:
        DatabaseManager.execute_query(
            "UPDATE requests SET status = 'rejected' WHERE request_id = %s",
            (request_id,),
            fetch=False,
            commit=True,
        )
        return True
    except Exception as e:
        logging.error(f"Failed to reject request: {str(e)}")
        st.error(f"Failed to reject request: {str(e)}")
        return False


def validate_file_upload(
    uploaded_file, file_type="pdf", max_size_mb=10, allowed_types=None
):
    """Validate file uploads for security and size constraints"""
    if allowed_types is None:
        allowed_types = {
            "pdf": ["application/pdf"],
            "image": ["image/jpeg", "image/png", "image/gif"],
            "document": [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ],
        }

    if uploaded_file is None:
        return False, "No file uploaded"

    # Check file size (max_size_mb in megabytes)
    max_bytes = max_size_mb * 1024 * 1024
    if uploaded_file.size > max_bytes:
        return False, f"File size exceeds {max_size_mb}MB limit"

    # Check file extension
    file_ext = os.path.splitext(uploaded_file.name.lower())[1][
        1:
    ]  # Get extension without dot
    valid_extensions = {
        "pdf": ["pdf"],
        "image": ["jpg", "jpeg", "png", "gif"],
        "document": ["doc", "docx"],
    }

    if file_ext not in valid_extensions.get(file_type, []):
        return (
            False,
            f"Only {', '.join(valid_extensions.get(file_type, []))} files are allowed",
        )

    # Generate file hash for tracking
    file_hash = hashlib.md5(uploaded_file.getvalue()).hexdigest()

    # Log file upload attempt
    logging.info(
        f"File upload: {uploaded_file.name}, Size: {uploaded_file.size}, Hash: {file_hash}"
    )

    return True, "Valid file"


# Usage in your form:
def file_upload_section():
    st.subheader("Supporting Documents")

    # Create columns for side-by-side upload fields
    col1, col2 = st.columns(2)

    with col1:
        acceptance_letter = st.file_uploader(
            "Upload Acceptance Letter",
            type=["pdf"],
            help="Upload acceptance letter as PDF (max 10MB)",
        )
        if acceptance_letter:
            is_valid, message = validate_file_upload(acceptance_letter)
            if not is_valid:
                st.error(message)
            else:
                st.success("✓ Acceptance letter uploaded successfully")
                st.info(
                    f"Filename: {acceptance_letter.name}, Size: {round(acceptance_letter.size/1024, 1)}KB"
                )

    with col2:
        research_paper = st.file_uploader(
            "Upload Research Paper",
            type=["pdf"],
            help="Upload research paper as PDF (max 10MB)",
        )
        if research_paper:
            is_valid, message = validate_file_upload(research_paper)
            if not is_valid:
                st.error(message)
            else:
                st.success("✓ Research paper uploaded successfully")
                st.info(
                    f"Filename: {research_paper.name}, Size: {round(research_paper.size/1024, 1)}KB"
                )

    return acceptance_letter, research_paper


def validate_date_range(
    date_from, date_to, min_days=1, max_days=30, future_days_required=5
):
    """Validate date range with comprehensive checks"""
    if date_from is None or date_to is None:
        return False, "Both start and end dates must be provided"

    # Convert to datetime.date if they aren't already
    if isinstance(date_from, str):
        try:
            date_from = datetime.strptime(date_from, "%Y-%m-%d").date()
        except ValueError:
            return False, "Start date format is incorrect. Please use YYYY-MM-DD."

    if isinstance(date_to, str):
        try:
            date_to = datetime.strptime(date_to, "%Y-%m-%d").date()
        except ValueError:
            return False, "End date format is incorrect. Please use YYYY-MM-DD."

    # Basic sequence check
    if date_from > date_to:
        return False, "Start date cannot be after end date"

    # Check against past dates
    today = datetime.now().date()
    if date_from < today:
        return False, "Travel cannot start in the past"

    # Check advance booking policy
    if (date_from - today).days < future_days_required:
        return (
            False,
            f"Requests must be made at least {future_days_required} days in advance",
        )

    # Check travel duration limits
    days_difference = (date_to - date_from).days + 1
    if days_difference < min_days:
        return False, f"Travel duration must be at least {min_days} day(s)"

    if days_difference > max_days:
        return False, f"Travel duration cannot exceed {max_days} days"

    # Check for restricted dates
    restricted_dates = get_restricted_dates()
    for period in restricted_dates:
        if date_from <= period["end_date"] and date_to >= period["start_date"]:
            return (
                False,
                f"Travel dates overlap with restricted period: {period['description']} ({period['start_date']} to {period['end_date']})",
            )

    return True, None


def get_restricted_dates():
    """Get all restricted date periods from the database"""
    try:
        result = DatabaseManager.execute_query(
            "SELECT start_date, end_date, description FROM restricteddates WHERE end_date >= CURRENT_DATE()"
        )
        return result
    except Exception as e:
        logging.error(f"Error fetching restricted dates: {str(e)}")
        return []


# Usage in your travel form:
def travel_dates_section():
    st.subheader("Travel Dates")

    col1, col2 = st.columns(2)
    with col1:
        date_from = st.date_input(
            "From Date",
            min_value=datetime.now().date(),
            help="Select your travel start date",
        )

    with col2:
        date_to = st.date_input(
            "To Date",
            min_value=date_from if date_from else datetime.now().date(),
            help="Select your travel end date",
        )

    if date_from and date_to:
        is_valid, message = validate_date_range(
            date_from, date_to, min_days=1, max_days=30, future_days_required=5
        )

        if not is_valid:
            st.error(message)
            st.session_state.dates_valid = False
        else:
            st.success("✓ Travel dates are valid")
            st.session_state.dates_valid = True

            # Show duration calculation
            duration = (date_to - date_from).days + 1
            st.info(f"Total trip duration: {duration} days")

    return date_from, date_to


# Replace the dynamic ORDER BY construction
def search_travel_requests(
    faculty_id=None,
    status=None,
    date_from=None,
    date_to=None,
    destination=None,
    limit=100,
    offset=0,
    sort_by="submission_date",
    sort_order="DESC",
):
    """Search travel requests with multiple optional filters using parameterized queries"""
    # Validate sort_by to prevent SQL injection
    allowed_sort_fields = [
        "request_id",
        "conference_name",
        "status",
        "submission_date",
        "date_from",
        "date_to",
        "destination",
        "city",
        "total_cost",
    ]

    # Default to submission_date if sort_by is not in allowed fields
    if sort_by not in allowed_sort_fields:
        sort_by = "submission_date"

    # Validate sort_order
    sort_order = "DESC" if sort_order.upper() != "ASC" else "ASC"

    # Start with base query
    query = """
        SELECT r.request_id, r.conference_name, r.status, r.submission_date, 
               r.date_from, r.date_to, r.destination, r.city,
               r.per_diem + r.registration_fee + r.visa_fee as total_cost,
               f.name as faculty_name
        FROM requests r
        JOIN faculty f ON r.faculty_user_id = f.user_id
        WHERE 1=1
    """

    # Initialize parameters list
    params = []

    # Add conditional filters with parameterized queries
    if faculty_id:
        query += " AND r.faculty_user_id = %s"
        params.append(faculty_id)

    if status:
        query += " AND r.status = %s"
        params.append(status)

    if date_from:
        query += " AND r.date_from >= %s"
        params.append(date_from)

    if date_to:
        query += " AND r.date_to <= %s"
        params.append(date_to)

    if destination:
        query += " AND r.destination LIKE %s"
        params.append(f"%{destination}%")

    # Add validated sorting
    query += f" ORDER BY r.{sort_by} {sort_order}"

    # Add pagination
    query += " LIMIT %s OFFSET %s"
    params.extend([limit, offset])

    # Execute query
    try:
        results = DatabaseManager.execute_query(query, tuple(params))
        return results
    except Exception as e:
        logging.error(f"Error in search_travel_requests: {str(e)}")
        return []


# Usage in your application
def display_search_interface():
    st.subheader("Search Travel Requests")

    # Create filter area
    with st.expander("Search Filters", expanded=True):
        col1, col2 = st.columns(2)

        with col1:
            status_filter = st.selectbox(
                "Status", options=["All", "pending", "approved", "rejected"], index=0
            )

            destination_filter = st.text_input("Destination", value="")

        with col2:
            start_date = st.date_input("From Date", value=None)
            end_date = st.date_input("To Date", value=None)

        # Search button
        if st.button("Search Requests"):
            # Convert "All" to None for the query
            status = None if status_filter == "All" else status_filter

            # Get results
            results = search_travel_requests(
                faculty_id=None,  # None means all faculty
                status=status,
                date_from=start_date,
                date_to=end_date,
                destination=destination_filter if destination_filter else None,
            )

            if results:
                st.success(f"Found {len(results)} matching requests")

                # Convert to DataFrame for display
                df = pd.DataFrame(results)
                st.dataframe(df, use_container_width=True)
            else:
                st.info("No matching requests found")


def setup_session_state():
    """Initialize session state variables"""
    # Basic session variables
    if "logged_in_user" not in st.session_state:
        st.session_state.logged_in_user = None
    if "user_role" not in st.session_state:
        st.session_state.user_role = None
    if "page" not in st.session_state:
        st.session_state.page = "login"

    # Security-related session variables
    if "last_activity" not in st.session_state:
        st.session_state.last_activity = datetime.now()
    if "login_attempts" not in st.session_state:
        st.session_state.login_attempts = 0
    if "lockout_until" not in st.session_state:
        st.session_state.lockout_until = None
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())


def check_session_security():
    """Check for session timeout and other security conditions"""
    # Check for account lockout
    if st.session_state.get("lockout_until"):
        if datetime.now() < st.session_state.lockout_until:
            remaining = st.session_state.lockout_until - datetime.now()
            st.error(
                f"Account temporarily locked due to multiple failed login attempts. Try again in {remaining.seconds//60} minutes."
            )
            return False
        else:
            st.session_state.lockout_until = None
            st.session_state.login_attempts = 0

    # Check for session timeout if user is logged in
    if st.session_state.logged_in_user:
        timeout_minutes = 30  # Session timeout after 30 minutes of inactivity
        current_time = datetime.now()
        time_diff = current_time - st.session_state.last_activity

        if time_diff.total_seconds() > (timeout_minutes * 60):
            # Session expired
            logout()
            st.warning(
                f"Your session has expired after {timeout_minutes} minutes of inactivity. Please log in again."
            )
            return False

        # Update last activity time
        st.session_state.last_activity = current_time

        # Log periodic session activity
        if time_diff.total_seconds() > 300:  # Log every 5 minutes
            logging.info(
                f"Active session: {st.session_state.session_id} for user {st.session_state.logged_in_user}"
            )

    return True


def secure_login(user_id, password):
    """Enhanced secure login with rate limiting and logging"""
    # Check for account lockout
    if st.session_state.get("lockout_until"):
        if datetime.now() < st.session_state.lockout_until:
            remaining = st.session_state.lockout_until - datetime.now()
            st.error(
                f"Account temporarily locked. Try again in {remaining.seconds//60} minutes."
            )
            return False

    # Basic input validation
    if not user_id or not password:
        st.error("Please provide both ID and password.")
        return False

    # Check credentials
    hashed_password = hashlib.sha256(password.encode()).hexdigest()

    try:
        result = DatabaseManager.execute_query(
            "SELECT * FROM faculty WHERE user_id = %s AND password = %s",
            (user_id, hashed_password),
        )

        if result and len(result) > 0:
            # Reset login attempts on successful login
            st.session_state.login_attempts = 0

            # Set user session data
            st.session_state.logged_in_user = user_id
            st.session_state.user_role = result[0]["role"]
            st.session_state.user_name = result[0]["name"]
            st.session_state.page = "main"
            st.session_state.last_activity = datetime.now()

            # Log successful login
            ip_address = get_client_ip()
            logging.info(f"Login successful: {user_id} from IP {ip_address}")

            # Record login in activity log
            record_user_activity(
                user_id,
                "login",
                {"ip_address": ip_address, "session_id": st.session_state.session_id},
            )

            return True
        else:
            # Increment failed login attempts
            st.session_state.login_attempts += 1

            # Log failed login attempt
            ip_address = get_client_ip()
            logging.warning(
                f"Failed login attempt: {user_id} from IP {ip_address}. Attempt #{st.session_state.login_attempts}"
            )

            # Implement account lockout after multiple failed attempts
            if st.session_state.login_attempts >= 5:
                lockout_time = datetime.now() + timedelta(minutes=15)
                st.session_state.lockout_until = lockout_time
                logging.warning(
                    f"Account {user_id} locked until {lockout_time} due to multiple failed login attempts."
                )
                st.error(
                    "Too many failed login attempts. Your account has been temporarily locked for 15 minutes."
                )
            else:
                st.error("Invalid ID or password.")

            return False

    except Exception as e:
        log_error(e, {"action": "login", "user_id": user_id})
        st.error("An error occurred during login. Please try again.")
        return False


def get_client_ip():
    """Get the client's IP address (placeholder)"""
    # In a production environment, this would extract the real IP
    # For Streamlit Cloud, you might use different headers
    return "127.0.0.1"  # Placeholder


def record_user_activity(user_id, activity_type, details=None):
    """Record user activity for audit purposes"""
    try:
        details_json = json.dumps(details) if details else "{}"

        DatabaseManager.execute_query(
            """
            INSERT INTO user_activity_log 
            (user_id, activity_type, details, ip_address, timestamp)
            VALUES (%s, %s, %s, %s, NOW())
            """,
            (user_id, activity_type, details_json, get_client_ip()),
            fetch=False,
            commit=True,
        )
    except Exception as e:
        logging.error(f"Error recording user activity: {str(e)}")


# Usage in your main application flow
def main():
    # Initialize session state
    setup_session_state()

    # Check session security
    if not check_session_security():
        # If session timeout occurred, don't proceed further
        st.stop()

    # Main application flow
    if st.session_state.page == "main" and st.session_state.logged_in_user:
        if st.session_state.get("show_success"):
            st.success("Login successful!")
            st.session_state.show_success = False
        main_app()
    else:
        login_page()


if __name__ == "__main__":
    setup_logging()

    # Initialize session state variables
    setup_session_state()

    # Check session security and proceed to main application flow
    if check_session_security():
        # Process logout request if needed
        if st.session_state.get("logout_requested", False):
            st.session_state.logout_requested = False
            st.rerun()

        # Display login page if not logged in, else show main app
        if st.session_state.page == "main" and st.session_state.logged_in_user:
            if st.session_state.get("show_success"):
                st.success("Login successful!")
                st.session_state.show_success = False
            main_app()
        else:
            login_page()


def display_activity_dashboard():
    """Display user activity dashboard for administrators"""
    st.title("User Activity Dashboard")

    # Date range selector
    col1, col2 = st.columns(2)
    with col1:
        start_date = st.date_input(
            "From Date", value=datetime.now().date() - timedelta(days=30)
        )
    with col2:
        end_date = st.date_input("To Date", value=datetime.now().date())

    # Fetch activity data
    activities = fetch_user_activities(start_date, end_date)

    if not activities:
        st.info("No activity data found for the selected period.")
        return

    # Convert to DataFrame for analysis
    df = pd.DataFrame(activities)
    df["date"] = pd.to_datetime(df["timestamp"]).dt.date

    # Display metrics
    st.subheader("Activity Metrics")
    metrics_col1, metrics_col2, metrics_col3 = st.columns(3)

    with metrics_col1:
        total_logins = len(df[df["activity_type"] == "login"])
        st.metric("Total Logins", total_logins)

    with metrics_col2:
        unique_users = df["user_id"].nunique()
        st.metric("Unique Users", unique_users)

    with metrics_col3:
        if "failed_login" in df["activity_type"].values:
            failed_logins = len(df[df["activity_type"] == "failed_login"])
            st.metric("Failed Login Attempts", failed_logins)
        else:
            st.metric("Failed Login Attempts", 0)

    # Activity by date chart
    st.subheader("Activity Timeline")
    daily_activity = (
        df.groupby(["date", "activity_type"]).size().reset_index(name="count")
    )

    fig = px.line(
        daily_activity,
        x="date",
        y="count",
        color="activity_type",
        title="User Activity by Date",
        labels={
            "date": "Date",
            "count": "Number of Activities",
            "activity_type": "Activity Type",
        },
    )
    st.plotly_chart(fig, use_container_width=True)

    # Activity breakdown
    st.subheader("Activity Breakdown")
    activity_counts = df["activity_type"].value_counts().reset_index()
    activity_counts.columns = ["Activity Type", "Count"]

    fig2 = px.pie(
        activity_counts,
        values="Count",
        names="Activity Type",
        title="Activity Distribution",
    )
    st.plotly_chart(fig2, use_container_width=True)

    # User activity table
    st.subheader("Recent Activity Log")
    recent_activities = df.sort_values("timestamp", ascending=False).head(100)
    st.dataframe(
        recent_activities[["timestamp", "user_id", "activity_type", "ip_address"]],
        use_container_width=True,
    )


def fetch_user_activities(start_date, end_date):
    """Fetch user activity data from the database"""
    try:
        result = DatabaseManager.execute_query(
            """
            SELECT * FROM user_activity_log
            WHERE DATE(timestamp) BETWEEN %s AND %s
            ORDER BY timestamp DESC
            """,
            (start_date, end_date),
        )
        return result
    except Exception as e:
        logging.error(f"Error fetching user activities: {str(e)}")
        return []
